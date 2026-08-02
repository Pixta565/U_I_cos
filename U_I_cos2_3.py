#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Программа анализа гармоник из CSV-файлов осциллографа (Rigol DHO814 и других)
в соответствии с ГОСТ 32144-2013 для напряжения 0.38 кВ.
Реализованы режимы: только напряжение, только ток, одна фаза (U+I),
трёхфазный объединённый (3 файла). Длительный анализ, настройки масштабов,
допусков, выбор папки отчётов, логирование, графическое отображение.
Сценарии 1, 2, 3: расширенные отчёты Word, Excel, фото для режима "Только напряжение".
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import csv
import os
import re
import threading
import queue
import numpy as np
from numpy.fft import rfft, rfftfreq
from datetime import datetime

# Графика
try:
    import matplotlib
    matplotlib.use("TkAgg")
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    HAVE_MPL = True
except ImportError:
    HAVE_MPL = False

# Word
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

# Excel
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from openpyxl.chart import BarChart, Reference
    from openpyxl.utils import get_column_letter
    HAVE_OPENPYXL = True
except ImportError:
    HAVE_OPENPYXL = False

# -------------------- ГОСТ 32144-2013 (0.38 кВ) --------------------
GOST_THD_LIMITS = {"norm": 8.0, "max": 12.0}
GOST_FREQ_DEV_NORM = 0.2   # Гц
GOST_FREQ_DEV_MAX = 0.4    # Гц
GOST_VOLTAGE_DEV = 3.0     # % (допуск RMS по умолчанию, но в настройках может меняться)

GOST_HARM_NON_MULT3 = {
    5: 6.0, 7: 5.0, 11: 3.5, 13: 3.0,
    17: 2.0, 19: 1.5, 23: 1.5, 25: 1.5,
    29: 1.0, 31: 1.0, 35: 0.7, 37: 0.7
}
GOST_HARM_MULT3 = {
    3: 5.0, 9: 1.5, 15: 0.3, 21: 0.2,
    27: 0.2, 33: 0.2, 39: 0.2
}
GOST_HARM_EVEN = {
    2: 2.0, 4: 1.0, 6: 0.5, 8: 0.5,
    10: 0.5, 12: 0.5, 14: 0.5, 16: 0.5,
    18: 0.5, 20: 0.5, 22: 0.5, 24: 0.5,
    26: 0.5, 28: 0.5, 30: 0.5, 32: 0.5,
    34: 0.5, 36: 0.5, 38: 0.5, 40: 0.5
}
GOST_HARM_ALL = {}
GOST_HARM_ALL.update(GOST_HARM_NON_MULT3)
GOST_HARM_ALL.update(GOST_HARM_MULT3)
GOST_HARM_ALL.update(GOST_HARM_EVEN)

# -------------------- Вспомогательные функции --------------------
def parse_csv(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        lines = list(reader)

    if len(lines) < 2:
        raise ValueError("Файл CSV слишком короткий")

    # Ищем t0 и tInc в первых двух строках
    t0_val = None
    tInc_val = None
    data_start = 1  # по умолчанию данные начинаются со второй строки

    # Сначала пробуем первую строку
    for field in lines[0]:
        field = field.strip()
        if field.startswith('t0 =') or field.startswith('t0='):
            t0_str = field.split('=', 1)[1].strip()
            try:
                t0_val = float(t0_str)
            except ValueError:
                pass
        elif field.startswith('tInc =') or field.startswith('tInc='):
            tInc_str = field.split('=', 1)[1].strip()
            try:
                tInc_val = float(tInc_str)
            except ValueError:
                pass

    # Если не нашли в первой, ищем во второй строке (старый формат)
    if t0_val is None or tInc_val is None:
        if len(lines) >= 2:
            second_line = lines[1]
            for field in second_line:
                field = field.strip()
                if (field.startswith('t0 =') or field.startswith('t0=')) and t0_val is None:
                    t0_val = float(field.split('=', 1)[1].strip())
                elif (field.startswith('tInc =') or field.startswith('tInc=')) and tInc_val is None:
                    tInc_val = float(field.split('=', 1)[1].strip())
            if t0_val is not None and tInc_val is not None:
                data_start = 2  # данные со третьей строки

    if t0_val is None or tInc_val is None:
        raise ValueError("Не удалось извлечь t0 и tInc")

    # Собираем имена каналов из первой строки (всё, что не является временными метками)
    ch_names = []
    for field in lines[0]:
        field = field.strip()
        if field == '' or field.startswith('t0') or field.startswith('tInc'):
            continue
        ch_names.append(field)

    # Чтение числовых данных
    data_rows = []
    for row in lines[data_start:]:
        cleaned = [v for v in row if v.strip() != '']
        if not cleaned:
            continue
        try:
            nums = [float(v) for v in cleaned]
        except ValueError:
            continue
        # Проверяем, что количество числовых столбцов соответствует числу каналов
        if len(nums) >= len(ch_names):
            data_rows.append(nums[:len(ch_names)])
        else:
            # Если в строке меньше чисел, чем каналов, дополним NaN или пропустим
            continue

    if not data_rows:
        raise ValueError("Нет данных для анализа")

    data = np.array(data_rows)
    return t0_val, tInc_val, ch_names, data

def estimate_frequency(signal, tInc, nominal=50.0, search_range=(45.0, 55.0)):
    N = len(signal)
    window = np.hanning(N)
    y = signal * window
    Y = rfft(y)
    freqs = rfftfreq(N, tInc)
    mask = (freqs >= search_range[0]) & (freqs <= search_range[1])
    if not np.any(mask):
        return nominal
    peak_idx = np.argmax(np.abs(Y[mask])) + np.where(mask)[0][0]
    freq_est = freqs[peak_idx]
    if 0 < peak_idx < len(freqs)-1:
        alpha = np.abs(Y[peak_idx-1])
        beta = np.abs(Y[peak_idx])
        gamma = np.abs(Y[peak_idx+1])
        if (alpha + gamma - 2*beta) != 0:
            delta = 0.5 * (alpha - gamma) / (alpha + gamma - 2*beta)
            freq_est = (peak_idx + delta) * (1.0/(N*tInc))
    return freq_est

def get_harmonic_amplitudes(signal, tInc, f0, max_harm=40):
    N = len(signal)
    window = np.hanning(N)
    y = signal * window
    Y = rfft(y)
    freqs = rfftfreq(N, tInc)
    df = 1.0/(N*tInc)
    harmonics = {}
    for h in range(1, max_harm+1):
        f_h = h * f0
        k = int(round(f_h / df))
        if k <= 0 or k >= len(freqs)-1:
            continue
        amp_peak = 4.0 * np.abs(Y[k]) / N
        harmonics[h] = (amp_peak, Y[k])
    return harmonics

def compute_harmonic_rms(harmonics):
    return {h: amp_peak/np.sqrt(2) for h, (amp_peak, _) in harmonics.items()}

def calculate_thd(harmonic_rms, fundamental_rms):
    sum_sq = sum(v**2 for h, v in harmonic_rms.items() if h >= 2)
    if fundamental_rms == 0:
        return 0.0
    return np.sqrt(sum_sq) / fundamental_rms * 100

def power_from_spectra(harm_v, harm_i, f0, max_harm=40):
    P_total = 0.0; Q_total = 0.0
    P_harm = {}; Q_harm = {}
    for h in range(1, max_harm+1):
        if h not in harm_v or h not in harm_i:
            continue
        _, Vc = harm_v[h]
        _, Ic = harm_i[h]
        S_complex = Vc * np.conj(Ic) / 2.0
        P = np.real(S_complex); Q = np.imag(S_complex)
        P_harm[h] = P; Q_harm[h] = Q
        P_total += P; Q_total += Q
    Vrms_total = np.sqrt(sum((amp/np.sqrt(2))**2 for amp, _ in harm_v.values()))
    Irms_total = np.sqrt(sum((amp/np.sqrt(2))**2 for amp, _ in harm_i.values()))
    S_total = Vrms_total * Irms_total
    pf = P_total / S_total if S_total != 0 else 0.0
    return {'P_total': P_total, 'Q_total': Q_total, 'S_total': S_total,
            'pf': pf, 'P_harm': P_harm, 'Q_harm': Q_harm}

def check_gost_limits(voltage_rms_harm, nominal_voltage=230):
    violations = []
    if 1 not in voltage_rms_harm:
        return violations
    V1 = voltage_rms_harm[1]
    for h, v_rms in voltage_rms_harm.items():
        if h == 1: continue
        if h in GOST_HARM_ALL:
            limit_percent = GOST_HARM_ALL[h]
            limit_val = (limit_percent / 100) * V1
            if v_rms > limit_val:
                violations.append(
                    f"Гармоника {h}: {v_rms:.3f} В > {limit_val:.3f} В ({limit_percent}% от U1)")
    thd_val = calculate_thd(voltage_rms_harm, V1)
    if thd_val > GOST_THD_LIMITS['max']:
        violations.append(f"THD: {thd_val:.2f}% > предельно допустимого {GOST_THD_LIMITS['max']}%")
    elif thd_val > GOST_THD_LIMITS['norm']:
        violations.append(f"THD: {thd_val:.2f}% > нормально допустимого {GOST_THD_LIMITS['norm']}%")
    return violations

def compute_period_rms(signal, tInc, f0):
    """Вычисление RMS для каждого периода основной частоты."""
    T = 1.0 / f0
    N_period = int(round(T / tInc))
    if N_period < 2:
        N_period = 2
    rms_values = []
    times = []
    start = 0
    while start + N_period <= len(signal):
        chunk = signal[start:start+N_period]
        rms = np.sqrt(np.mean(chunk**2))
        t_center = (start + N_period/2) * tInc  # время середины периода
        rms_values.append(rms)
        times.append(t_center)
        start += N_period
    return times, rms_values

def compute_signal_stats(signal, tInc, f0, nominal_voltage, rms_tolerance=3.0):
    """
    Возвращает словарь с основными параметрами сигнала напряжения.
    """
    N = len(signal)
    t = np.arange(N) * tInc
    # Основные параметры
    Urms_total = np.sqrt(np.mean(signal**2))
    Upeak = np.max(np.abs(signal))
    Umax_inst = np.max(signal)
    Umin_inst = np.min(signal)
    crest_factor = Upeak / Urms_total if Urms_total != 0 else 0.0

    # Отклонение напряжения
    deviation = (Urms_total - nominal_voltage) / nominal_voltage * 100

    # Для длительного анализа: RMS по периодам
    times_period, rms_periods = compute_period_rms(signal, tInc, f0)
    if rms_periods:
        max_rms_period = np.max(rms_periods)
        min_rms_period = np.min(rms_periods)
        avg_rms_period = np.mean(rms_periods)
    else:
        max_rms_period = min_rms_period = avg_rms_period = Urms_total

    return {
        'Urms': Urms_total,
        'Upeak': Upeak,
        'Umax_inst': Umax_inst,
        'Umin_inst': Umin_inst,
        'crest_factor': crest_factor,
        'deviation_percent': deviation,
        'f0': f0,
        'rms_periods': rms_periods,
        'times_period': times_period,
        'max_rms_period': max_rms_period,
        'min_rms_period': min_rms_period,
        'avg_rms_period': avg_rms_period,
        'tInc': tInc,
        'N_samples': N
    }

# -------------------- Класс приложения --------------------
class HarmonicsApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Анализатор гармоник (ГОСТ 32144-2013, 0.38 кВ)")
        self.root.geometry("1200x800")
        self.style = ttk.Style()
        self.style.theme_use('clam')

        self.mode_var = tk.StringVar(value="Только напряжение")
        self.long_analysis_var = tk.BooleanVar(value=False)
        self.save_folder = tk.StringVar(value=os.getcwd())

        self.settings = {
            'scale_voltage': 1.0,
            'scale_current': 1694.0,
            'nominal_voltage': 230.0,
            'rms_tolerance': 3.0
        }

        self.csv_file_paths = [tk.StringVar() for _ in range(3)]
        self.channel_vars = []
        self.current_file_data = None
        self.result_queue = queue.Queue()
        self.have_plots = HAVE_MPL

        if not HAVE_DOCX:
            self.log("Библиотека python-docx не установлена. Отчёты Word не будут формироваться.")
        if not HAVE_OPENPYXL:
            self.log("Библиотека openpyxl не установлена. Отчёты Excel не будут формироваться.")

        self.setup_ui()
        self.check_queue()

    def setup_ui(self):
        main_panel = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_panel.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        left_frame = ttk.Frame(main_panel, width=400)
        right_frame = ttk.Frame(main_panel, width=800)
        main_panel.add(left_frame, weight=0)
        main_panel.add(right_frame, weight=1)

        self.build_control_panel(left_frame)
        self.build_output_panel(right_frame)

    def build_control_panel(self, parent):
        ttk.Label(parent, text="Параметры анализа", font=('Arial', 11, 'bold')).pack(anchor=tk.W, pady=5)

        mode_frame = ttk.LabelFrame(parent, text="Режим анализа", padding=5)
        mode_frame.pack(fill=tk.X, pady=5)
        modes = [
            "Только напряжение",
            "Только ток",
            "Напряжение и ток (одна фаза)",
            "Трёхфазный объединённый (3 файла)"
        ]
        mode_combo = ttk.Combobox(mode_frame, textvariable=self.mode_var, values=modes, state="readonly")
        mode_combo.pack(fill=tk.X, padx=5, pady=2)
        mode_combo.bind('<<ComboboxSelected>>', self.on_mode_change)

        self.file_frame = ttk.LabelFrame(parent, text="Файлы CSV", padding=5)
        self.file_frame.pack(fill=tk.X, pady=5)
        self.file_widgets = []
        for i in range(3):
            row = ttk.Frame(self.file_frame)
            row.pack(fill=tk.X, pady=2)
            lbl = ttk.Label(row, text=f"Файл {i+1}:", width=8)
            lbl.pack(side=tk.LEFT)
            entry = ttk.Entry(row, textvariable=self.csv_file_paths[i], state="readonly")
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
            btn = ttk.Button(row, text="...", width=3, command=lambda idx=i: self.browse_file(idx))
            btn.pack(side=tk.LEFT)
            self.file_widgets.append((row, lbl, entry, btn))
        self.update_file_visibility()

        long_check = ttk.Checkbutton(parent, text="Длительный анализ (большие файлы)", variable=self.long_analysis_var)
        long_check.pack(anchor=tk.W, pady=5)

        chan_frame = ttk.LabelFrame(parent, text="Назначение каналов", padding=5)
        chan_frame.pack(fill=tk.X, pady=5)
        self.chan_assignment_frame = ttk.Frame(chan_frame)
        self.chan_assignment_frame.pack(fill=tk.X)

        save_frame = ttk.Frame(parent)
        save_frame.pack(fill=tk.X, pady=5)
        ttk.Label(save_frame, text="Папка отчётов:").pack(side=tk.LEFT)
        ttk.Entry(save_frame, textvariable=self.save_folder, state="readonly", width=20).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        ttk.Button(save_frame, text="...", width=3, command=self.browse_save_folder).pack(side=tk.LEFT)

        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, pady=10)
        ttk.Button(btn_frame, text="Настройки", command=self.open_settings).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Запуск", command=self.start_analysis).pack(side=tk.RIGHT, padx=2)

        self.progress = ttk.Progressbar(parent, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=5)

    def build_output_panel(self, parent):
        if self.have_plots:
            plot_frame = ttk.LabelFrame(parent, text="Графики", padding=5)
            plot_frame.pack(fill=tk.BOTH, expand=True, pady=5)
            self.fig = Figure(figsize=(6, 3), dpi=100)
            self.ax1 = self.fig.add_subplot(121)
            self.ax2 = self.fig.add_subplot(122)
            self.canvas = FigureCanvasTkAgg(self.fig, master=plot_frame)
            self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        else:
            ttk.Label(parent, text="Matplotlib не установлен - графики отключены").pack()

        log_frame = ttk.LabelFrame(parent, text="Логи / отчёт", padding=5)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.log_text = scrolledtext.ScrolledText(log_frame, height=10, wrap=tk.WORD)
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def update_file_visibility(self):
        mode = self.mode_var.get()
        num_files = 3 if "Трёхфазный" in mode else 1
        for i, widgets in enumerate(self.file_widgets):
            if i < num_files:
                widgets[0].pack(fill=tk.X, pady=2)
            else:
                widgets[0].pack_forget()

    def on_mode_change(self, event=None):
        self.update_file_visibility()
        self.clear_channel_assignment()
        self.channel_vars.clear()

    def clear_channel_assignment(self):
        for widget in self.chan_assignment_frame.winfo_children():
            widget.destroy()

    def browse_file(self, idx):
        filename = filedialog.askopenfilename(
            title="Выберите CSV файл",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if filename:
            self.csv_file_paths[idx].set(filename)
            self.update_channel_choices()

    def update_channel_choices(self):
        if "Трёхфазный" in self.mode_var.get():
            for i in range(3):
                path = self.csv_file_paths[i].get()
                if path:
                    break
            else:
                return
        else:
            path = self.csv_file_paths[0].get()
            if not path:
                return

        try:
            _, _, ch_names, _ = parse_csv(path)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл: {e}")
            return

        self.clear_channel_assignment()
        self.channel_vars = []
        options = ["Не используется", "Напряжение", "Ток"]
        for i, name in enumerate(ch_names):
            frame = ttk.Frame(self.chan_assignment_frame)
            frame.pack(fill=tk.X, pady=1)
            ttk.Label(frame, text=f"{name}:", width=10).pack(side=tk.LEFT)
            var = tk.StringVar(value="Не используется")
            combo = ttk.Combobox(frame, textvariable=var, values=options, state="readonly", width=15)
            combo.pack(side=tk.LEFT, padx=5)
            self.channel_vars.append(var)

    def browse_save_folder(self):
        folder = filedialog.askdirectory(title="Выберите папку для сохранения отчётов")
        if folder:
            self.save_folder.set(folder)

    def open_settings(self):
        settings_win = tk.Toplevel(self.root)
        settings_win.title("Настройки")
        settings_win.geometry("300x250")
        settings_win.resizable(False, False)

        ttk.Label(settings_win, text="Масштаб напряжения:").pack(pady=2)
        v_scale = tk.StringVar(value=str(self.settings['scale_voltage']))
        ttk.Entry(settings_win, textvariable=v_scale).pack(pady=2)

        ttk.Label(settings_win, text="Масштаб тока:").pack(pady=2)
        i_scale = tk.StringVar(value=str(self.settings['scale_current']))
        ttk.Entry(settings_win, textvariable=i_scale).pack(pady=2)

        ttk.Label(settings_win, text="Номинальное напряжение (В):").pack(pady=2)
        nom_v = tk.StringVar(value=str(self.settings['nominal_voltage']))
        ttk.Entry(settings_win, textvariable=nom_v).pack(pady=2)

        ttk.Label(settings_win, text="Допуск RMS (+- %):").pack(pady=2)
        tol = tk.StringVar(value=str(self.settings['rms_tolerance']))
        ttk.Entry(settings_win, textvariable=tol).pack(pady=2)

        def save_settings():
            try:
                self.settings['scale_voltage'] = float(v_scale.get())
                self.settings['scale_current'] = float(i_scale.get())
                self.settings['nominal_voltage'] = float(nom_v.get())
                self.settings['rms_tolerance'] = float(tol.get())
            except ValueError:
                messagebox.showerror("Ошибка", "Некорректное числовое значение")
                return
            settings_win.destroy()

        ttk.Button(settings_win, text="Сохранить", command=save_settings).pack(pady=10)

    def log(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def start_analysis(self):
        mode = self.mode_var.get()
        if "Трёхфазный" in mode:
            paths = [self.csv_file_paths[i].get() for i in range(3)]
            if not all(paths):
                messagebox.showerror("Ошибка", "Выберите три CSV файла")
                return
        else:
            if not self.csv_file_paths[0].get():
                messagebox.showerror("Ошибка", "Выберите CSV файл")
                return
            paths = [self.csv_file_paths[0].get()]

        if not self.channel_vars:
            messagebox.showerror("Ошибка", "Не выбраны назначения каналов. Загрузите файл и назначьте.")
            return

        self.progress.start()
        self.log("Начинаем анализ...")
        thread = threading.Thread(target=self.run_analysis, args=(mode, paths), daemon=True)
        thread.start()

    def run_analysis(self, mode, paths):
        try:
            if "Трёхфазный" in mode:
                self.analyze_three_phase(paths)
            else:
                self.analyze_single(mode, paths[0])
        except Exception as e:
            self.result_queue.put(("error", str(e)))
        finally:
            self.result_queue.put(("done", None))

    def analyze_single(self, mode, path):
        t0, tInc, ch_names, data = parse_csv(path)
        assignments = [var.get() for var in self.channel_vars]
        voltage_signals = []
        current_signals = []
        for i, assign in enumerate(assignments):
            if i >= data.shape[1]:
                continue
            signal = data[:, i]
            if assign == "Напряжение":
                voltage_signals.append(signal * self.settings['scale_voltage'])
            elif assign == "Ток":
                current_signals.append(signal * self.settings['scale_current'])

        if mode == "Только напряжение":
            if not voltage_signals:
                raise ValueError("Не выбраны каналы напряжения")
            self.analyze_voltage_only(voltage_signals, tInc, ch_names, assignments)
        elif mode == "Только ток":
            if not current_signals:
                raise ValueError("Не выбраны каналы тока")
            self.analyze_current_only(current_signals, tInc)
            results = {}
            for idx, sig in enumerate(current_signals):
                res = self.process_signal(sig, tInc, 'current')
                results[f"I{idx+1}"] = res
            self.result_queue.put(("current_results", results))
        elif mode == "Напряжение и ток (одна фаза)":
            if len(voltage_signals) != 1 or len(current_signals) != 1:
                raise ValueError("Назначьте ровно одно напряжение и один ток")
            v_sig, i_sig = voltage_signals[0], current_signals[0]
            res_v = self.process_signal(v_sig, tInc, 'voltage')
            res_i = self.process_signal(i_sig, tInc, 'current')
            harm_v, harm_i = self._get_harmonics(v_sig, i_sig, tInc, res_v['f0'])
            pow_dict = power_from_spectra(harm_v, harm_i, res_v['f0'])
            self.result_queue.put(("combined_results", {
                'voltage': res_v, 'current': res_i, 'power': pow_dict
            }))
        else:
            raise ValueError("Некорректная конфигурация каналов для выбранного режима")

    def analyze_voltage_only(self, voltage_signals, tInc, ch_names, assignments):
        """
        Расширенный анализ напряжения с детальными параметрами и таблицами.
        """
        long_analysis = self.long_analysis_var.get()
        nominal_voltage = self.settings['nominal_voltage']
        rms_tol = self.settings['rms_tolerance']

        phases = len(voltage_signals)
        phase_names = ['А', 'В', 'С'][:phases]

        # Сбор данных по каждой фазе
        phase_data = []
        for idx, sig in enumerate(voltage_signals):
            f0 = estimate_frequency(sig, tInc)
            stats = compute_signal_stats(sig, tInc, f0, nominal_voltage, rms_tol)
            harmonics = get_harmonic_amplitudes(sig, tInc, f0)
            rms_harm = compute_harmonic_rms(harmonics)
            thd = calculate_thd(rms_harm, stats['Urms'])
            violations = check_gost_limits(rms_harm, nominal_voltage)
            phase_data.append({
                'name': phase_names[idx],
                'signal': sig,
                'f0': f0,
                'stats': stats,
                'harmonics': harmonics,
                'rms_harm': rms_harm,
                'thd': thd,
                'violations': violations,
                'tInc': tInc
            })

        self.result_queue.put(("voltage_results_extended", {
            'phases': phase_data,
            'long_analysis': long_analysis,
            'tInc': tInc,
            'ch_names': ch_names,
            'assignments': assignments
        }))

    def analyze_current_only(self, current_signals, tInc):
        long_analysis = self.long_analysis_var.get()
        channels = []
        for idx, sig in enumerate(current_signals):
            # Удаление постоянной составляющей
            sig_ac = sig - np.mean(sig)
            f0 = estimate_frequency(sig_ac, tInc)
            harmonics = get_harmonic_amplitudes(sig_ac, tInc, f0)
            rms_harm = compute_harmonic_rms(harmonics)
            Irms = np.sqrt(np.mean(sig_ac**2))
            I1_rms = rms_harm.get(1, 0)
            thdi = calculate_thd(rms_harm, I1_rms) if I1_rms else 0.0
            I1_peak = harmonics.get(1, (0,))[0]
            Imax_inst = np.max(sig_ac)
            Imin_inst = np.min(sig_ac)
            Ipeak = np.max(np.abs(sig_ac))
            crest = Ipeak / Irms if Irms != 0 else 0.0
            times_period, rms_periods = compute_period_rms(sig_ac, tInc, f0)
            max_rms_period = np.max(rms_periods) if len(rms_periods) > 0 else Irms
            min_rms_period = np.min(rms_periods) if len(rms_periods) > 0 else Irms
            avg_rms_period = np.mean(rms_periods) if len(rms_periods) > 0 else Irms
            channels.append({
                'name': f'I{idx+1}',
                'signal': sig_ac,
                'f0': f0,
                'stats': {
                    'Irms': Irms,
                    'I1_rms': I1_rms,
                    'I1_peak': I1_peak,
                    'THDi': thdi,
                    'Imax_inst': Imax_inst,
                    'Imin_inst': Imin_inst,
                    'Ipeak': Ipeak,
                    'crest_factor': crest,
                    'rms_periods': rms_periods,
                    'times_period': times_period,
                    'max_rms_period': max_rms_period,
                    'min_rms_period': min_rms_period,
                    'avg_rms_period': avg_rms_period,
                    'tInc': tInc,
                    'N_samples': len(sig_ac)
                },
                'harmonics': harmonics,
                'rms_harm': rms_harm,
                'thdi': thdi,
                'f0': f0,
                'tInc': tInc
            })
        self.result_queue.put(("current_results_extended", {
            'channels': channels,
            'long_analysis': long_analysis,
            'tInc': tInc
        }))


    def analyze_three_phase(self, paths):
        # Существующий трёхфазный анализ оставляем как есть
        phase_results = []
        total_power = {'P_total': 0, 'Q_total': 0, 'S_total': 0, 'pf': None}
        all_v_thds = []
        all_i_thds = []
        for phase, path in enumerate(paths, 1):
            self.result_queue.put(("log", f"Обработка фазы {phase}"))
            t0, tInc, ch_names, data = parse_csv(path)
            assignments = [var.get() for var in self.channel_vars]
            v_sig = None; i_sig = None
            for i, assign in enumerate(assignments):
                if i >= data.shape[1]: continue
                if assign == "Напряжение":
                    v_sig = data[:, i] * self.settings['scale_voltage']
                elif assign == "Ток":
                    i_sig = data[:, i] * self.settings['scale_current']
            if v_sig is None or i_sig is None:
                raise ValueError(f"Фаза {phase}: не назначены напряжение и ток")
            res_v = self.process_signal(v_sig, tInc, 'voltage')
            res_i = self.process_signal(i_sig, tInc, 'current')
            harm_v, harm_i = self._get_harmonics(v_sig, i_sig, tInc, res_v['f0'])
            pow_dict = power_from_spectra(harm_v, harm_i, res_v['f0'])
            phase_results.append({
                'phase': phase, 'voltage': res_v, 'current': res_i, 'power': pow_dict
            })
            total_power['P_total'] += pow_dict['P_total']
            total_power['Q_total'] += pow_dict['Q_total']
            total_power['S_total'] += pow_dict['S_total']
            all_v_thds.append(res_v['thd'])
            all_i_thds.append(res_i['thd'])
        if total_power['S_total'] != 0:
            total_power['pf'] = total_power['P_total'] / total_power['S_total']
        overall_thd_v = max(all_v_thds) if all_v_thds else 0
        overall_thd_i = max(all_i_thds) if all_i_thds else 0
        self.result_queue.put(("three_phase_results", {
            'phases': phase_results, 'total_power': total_power,
            'overall_thd_v': overall_thd_v, 'overall_thd_i': overall_thd_i
        }))

    def process_signal(self, signal, tInc, sig_type):
        f0 = estimate_frequency(signal, tInc)
        harmonics = get_harmonic_amplitudes(signal, tInc, f0)
        rms_harm = compute_harmonic_rms(harmonics)
        fundamental_rms = rms_harm.get(1, 0)
        thd = calculate_thd(rms_harm, fundamental_rms)
        result = {
            'f0': f0,
            'fundamental_rms': fundamental_rms,
            'harmonic_rms': rms_harm,
            'thd': thd,
            'type': sig_type
        }
        if sig_type == 'voltage':
            result['gost_violations'] = check_gost_limits(rms_harm, self.settings['nominal_voltage'])
        return result

    def _get_harmonics(self, v_sig, i_sig, tInc, f0=None):
        if f0 is None:
            f0 = estimate_frequency(v_sig, tInc)
        harm_v = get_harmonic_amplitudes(v_sig, tInc, f0)
        harm_i = get_harmonic_amplitudes(i_sig, tInc, f0)
        return harm_v, harm_i

    def check_queue(self):
        try:
            while True:
                msg = self.result_queue.get_nowait()
                if msg[0] == "log":
                    self.log(msg[1])
                elif msg[0] == "error":
                    self.log(f"ОШИБКА: {msg[1]}")
                    messagebox.showerror("Ошибка", msg[1])
                    self.progress.stop()
                elif msg[0] == "done":
                    self.progress.stop()
                    self.log("Анализ завершён.")
                elif msg[0] == "voltage_results_extended":
                    self.handle_extended_voltage_results(msg[1])
                elif msg[0] == "combined_results":
                    self.display_combined_results(msg[1])
                elif msg[0] == "three_phase_results":
                    self.display_three_phase_results(msg[1])
                elif msg[0] == "current_results_extended":
                    self.handle_extended_current_results(msg[1])
        except queue.Empty:
            pass
        self.root.after(100, self.check_queue)

    # -------------------- Обработчики расширенного отчёта по напряжению --------------------
    def handle_extended_voltage_results(self, data):
        phases = data['phases']
        long_analysis = data['long_analysis']
        tInc = data['tInc']
        ch_names = data['ch_names']
        assignments = data['assignments']

        # Логирование
        self.log("=== Расширенный анализ напряжения ===")
        for ph in phases:
            self.log(f"Фаза {ph['name']}: f0={ph['f0']:.3f} Гц, U={ph['stats']['Urms']:.2f} В, THD={ph['thd']:.2f}%")

        # Генерация Word, Excel, графиков
        if HAVE_MPL:
                    self.generate_extended_voltage_plots(data)
        if HAVE_DOCX:
            self.generate_extended_voltage_word(data)
        if HAVE_OPENPYXL:
            self.generate_extended_voltage_excel(data)
        

        # Также можно оставить текстовый отчёт в логах
        # Сохраним общий текстовый файл
        report = self.build_text_report(data)
        self.log(report)
        self.save_text_report(report, "Отчет_напряжение.txt")

    def handle_extended_current_results(self, data):
        channels = data['channels']
        long_analysis = data['long_analysis']
        tInc = data['tInc']

        self.log("=== Расширенный анализ тока ===")
        for ch in channels:
            self.log(f"Канал {ch['name']}: f0={ch['f0']:.3f} Гц, Irms={ch['stats']['Irms']:.2f} А, THDi={ch['thdi']:.2f}%")

        if HAVE_MPL:
            self.generate_extended_current_plots(data)
        if HAVE_DOCX:
            self.generate_extended_current_word(data)
        if HAVE_OPENPYXL:
            self.generate_extended_current_excel(data)

    def build_text_report(self, data):
        phases = data['phases']
        long_analysis = data['long_analysis']
        rep = "=== Отчёт по напряжению (ГОСТ 32144-2013) ===\n"
        rep += f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
        rep += f"Количество фаз: {len(phases)}\n"
        for ph in phases:
            rep += f"\nФаза {ph['name']}:\n"
            s = ph['stats']
            rep += f"  f0={ph['f0']:.3f} Гц, U={s['Urms']:.2f} В, THD={ph['thd']:.2f}%\n"
        return rep

    def save_text_report(self, text, filename):
        try:
            path = os.path.join(self.save_folder.get(), filename)
            with open(path, 'w', encoding='utf-8') as f:
                f.write(text)
            self.log(f"Текстовый отчёт сохранён: {path}")
        except Exception as e:
            self.log(f"Ошибка сохранения: {e}")

    # ========== WORD отчёты (расширенные) ==========
    def generate_extended_voltage_word(self, data):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as docx_qn
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        phases = data['phases']
        long_analysis = data['long_analysis']
        tInc = data['tInc']
        ch_names = data['ch_names']
        assignments = data['assignments']

        doc = self._setup_word_document()

        # --- Убираем режим совместимости ---
        settings = doc.settings.element
        for child in settings:
            if child.tag == docx_qn('w:compat'):
                settings.remove(child)
                break

        self._add_heading_word(doc, "Отчёт по напряжению")
        self._add_paragraph_word(doc, "")

        # ---------- Общие сведения ----------
        self._add_heading_word(doc, "Общие сведения", level=2)
        self._add_paragraph_word(doc, f"Дата и время: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
        num_ch = len([a for a in assignments if a != "Не используется"])
        self._add_paragraph_word(doc, f"Формат: Rigol DHO814 ({num_ch} канал(ов))")
        self._add_paragraph_word(doc, f"Частота дискретизации: {1.0/tInc:.2f} Гц")
        N_samples = len(phases[0]['signal']) if phases else 0
        self._add_paragraph_word(doc, f"Количество отсчётов: {N_samples}")
        self._add_paragraph_word(doc, f"Масштаб напряжения: x{self.settings['scale_voltage']:.2f}")
        self._add_paragraph_word(doc, f"Масштаб тока: x{self.settings['scale_current']:.2f}")
        self._add_paragraph_word(doc, "Инверсия тока: нет")
        voltage_channels = []
        for i, a in enumerate(assignments):
            if a == "Напряжение":
                voltage_channels.append(ch_names[i] if i < len(ch_names) else f"CH{i+1}")
        self._add_paragraph_word(doc, f"Канал(ы) напряжения: {', '.join(voltage_channels)}")
        self._add_paragraph_word(doc, f"Количество фаз: {len(phases)}")

        # ---------- Частота ----------
        self._add_heading_word(doc, "Частота", level=2)
        f0_avg = np.mean([ph['f0'] for ph in phases])
        f0_dev = f0_avg - 50.0
        status_f = "НОРМАЛЬНО ДОПУСТИМОЕ" if abs(f0_dev) <= GOST_FREQ_DEV_NORM else \
                "ПРЕДЕЛЬНО ДОПУСТИМОЕ" if abs(f0_dev) <= GOST_FREQ_DEV_MAX else "НЕДОПУСТИМОЕ"
        self._add_paragraph_word(doc, f"Основная частота: {f0_avg:.3f} Гц")
        self._add_paragraph_word(doc, f"Отклонение частоты: {f0_dev:.3f} Гц ({status_f})")

        # ---------- Таблица 1. Параметры напряжения ----------
        self._add_heading_word(doc, "Таблица 1. Параметры напряжения", level=2)
        self._add_paragraph_word(doc, f"Номинальное напряжение: {self.settings['nominal_voltage']:.1f} В")
        headers1 = ["Параметр"]
        for ph in phases:
            headers1.append(f"Фаза {ph['name']}")
        if len(phases) > 1:
            headers1.append("Ср. по фазам")

        table1 = doc.add_table(rows=1, cols=len(headers1))
        table1.style = 'Table Grid'
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Заголовок
        for i, hdr in enumerate(headers1):
            cell = table1.rows[0].cells[i]
            cell.text = hdr
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)

        def add_table1_row(name, values, unit="", tol_abs=None, tol_pct=None):
            row_cells = table1.add_row().cells
            row_cells[0].text = name
            for i, val in enumerate(values):
                cell = row_cells[i+1]
                cell.text = f"{val:.2f}{unit}"
                is_bold = False
                if tol_abs is not None and abs(val) > tol_abs:
                    is_bold = True
                if tol_pct is not None and abs(val) > tol_pct:
                    is_bold = True
                if is_bold:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
            if len(phases) > 1:
                avg_val = np.mean(values)
                cell = row_cells[-1]
                cell.text = f"{avg_val:.2f}{unit}"
                if tol_abs is not None and abs(avg_val) > tol_abs:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
            # форматирование ячеек
            for cell in row_cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

        s_list = [ph['stats'] for ph in phases]
        add_table1_row("Действующее напр. (RMS)", [s['Urms'] for s in s_list], " В")
        add_table1_row("Отклонение напряжения (±3%)", [s['deviation_percent'] for s in s_list], " %", tol_abs=3.0)
        add_table1_row("Макс. мгновенное", [s['Umax_inst'] for s in s_list], " В")
        add_table1_row("Мин. мгновенное", [s['Umin_inst'] for s in s_list], " В")
        add_table1_row("Пиковое напряжение", [s['Upeak'] for s in s_list], " В")
        add_table1_row("Коэф. амплитуды (норма ≤1.41)", [s['crest_factor'] for s in s_list], "", tol_abs=1.41)
        add_table1_row("Макс. RMS за период", [s['max_rms_period'] for s in s_list], " В")
        add_table1_row("Мин. RMS за период", [s['min_rms_period'] for s in s_list], " В")
        add_table1_row("Среднее RMS за период", [s['avg_rms_period'] for s in s_list], " В")

        self._format_table(table1)

        # ---------- Таблица 2. Параметры по фазам ----------
        self._add_heading_word(doc, "Таблица 2. Параметры по фазам", level=3)
        headers2 = ["Параметр"] + [f"Фаза {ph['name']}" for ph in phases] + (["Общее (среднее)"] if len(phases)>1 else [])
        table2 = doc.add_table(rows=1, cols=len(headers2))
        table2.style = 'Table Grid'
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hdr in enumerate(headers2):
            cell = table2.rows[0].cells[i]
            cell.text = hdr
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0,0,0)

        def add_table2_row(name, values, unit="", tol_check=None):
            row_cells = table2.add_row().cells
            row_cells[0].text = name
            for i, val in enumerate(values):
                cell = row_cells[i+1]
                cell.text = f"{val:.2f}{unit}"
                if tol_check and val > tol_check:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
            if len(phases) > 1:
                avg_val = np.mean(values)
                cell = row_cells[-1]
                cell.text = f"{avg_val:.2f}{unit}"
                if tol_check and avg_val > tol_check:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
            for cell in row_cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0,0,0)

        add_table2_row("f0, Гц", [ph['f0'] for ph in phases])
        add_table2_row("Urms, В", [ph['stats']['Urms'] for ph in phases])
        add_table2_row("U откл., %", [ph['stats']['deviation_percent'] for ph in phases])
        u1_peaks = [ph['harmonics'].get(1, (0,0))[0] for ph in phases]
        add_table2_row("Ампл. осн. гарм., В", u1_peaks)
        ku_values = [ph['thd'] for ph in phases]
        add_table2_row("Ku (THD), % (≤8%)", ku_values, tol_check=8.0)

        # Статус Ku
        row_status = table2.add_row().cells
        row_status[0].text = "Статус Ku"
        status_ku = []
        for thd in ku_values:
            if thd <= GOST_THD_LIMITS['norm']: status_ku.append("Норма")
            elif thd <= GOST_THD_LIMITS['max']: status_ku.append("Пред.")
            else: status_ku.append("Недоп.")
        if len(phases) > 1: status_ku.append("")
        for i, st in enumerate(status_ku):
            row_status[i+1].text = st

        # Нарушения и предупреждения
        viol_counts = [len(ph['violations']) for ph in phases]
        row_viol = table2.add_row().cells
        row_viol[0].text = "Нарушений"
        for i, v in enumerate(viol_counts):
            row_viol[i+1].text = str(v)
        if len(phases) > 1: row_viol[-1].text = ""

        row_warn = table2.add_row().cells
        row_warn[0].text = "Предупреждений"
        for i in range(len(phases)):
            row_warn[i+1].text = "0"
        if len(phases) > 1: row_warn[-1].text = ""

        # Применяем общее форматирование ко всей таблице
        for row in table2.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0,0,0)
        self._format_table(table2)

        # ---------- Таблица 3. Детальная таблица гармоник ----------
        self._add_heading_word(doc, "Таблица 3. Детальная таблица гармоник", level=3)
        max_harm = 40
        harm_table_header = ["№", "Частота, Гц"]
        for ph in phases:
            harm_table_header.append(f"Ампл. {ph['name']}, %")
        harm_table_header += ["Предел ГОСТ, %", "Отклонение, %", "Статус", "Таблица ГОСТ"]
        table3 = doc.add_table(rows=1, cols=len(harm_table_header))
        table3.style = 'Table Grid'
        table3.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hdr in enumerate(harm_table_header):
            cell = table3.rows[0].cells[i]
            cell.text = hdr
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0,0,0)

        f0_avg = np.mean([ph['f0'] for ph in phases])
        for h in range(1, max_harm+1):
            freq = h * f0_avg
            row_cells = table3.add_row().cells
            row_cells[0].text = str(h)
            row_cells[1].text = f"{freq:.1f}"
            col = 2
            rel_amps = []
            for ph in phases:
                U1_rms = ph['rms_harm'].get(1, ph['stats']['Urms'])
                rms_h = ph['rms_harm'].get(h, 0)
                rel = (rms_h / U1_rms * 100) if U1_rms else 0
                rel_amps.append(rel)
                row_cells[col].text = f"{rel:.2f}"
                col += 1
            limit = GOST_HARM_ALL.get(h, 0)
            row_cells[col].text = f"{limit:.2f}"; col+=1
            deviation = np.mean(rel_amps) - limit
            row_cells[col].text = f"{deviation:.2f}"; col+=1
            status = "Норма"
            for rel in rel_amps:
                if rel > limit and limit > 0:
                    status = "Превышение"
                    break
            row_cells[col].text = status; col+=1
            if h % 2 == 0: table_gost = "2 (чётные)"
            elif h % 3 == 0: table_gost = "4 (кратные 3)"
            else: table_gost = "3 (некратные 3)"
            row_cells[col].text = table_gost

            # Если превышение – выделяем всю строку жирным
            if status == "Превышение":
                for cell in row_cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

            # Форматирование ячеек
            for cell in row_cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0,0,0)

        self._format_table(table3)

        # ---------- Таблица 4 (если длительный анализ) ----------
        if long_analysis:
            self._add_heading_word(doc, "Таблица 4. Значения RMS напряжений для каждой фазы", level=3)
            ref_ph = phases[0]
            n_periods = len(ref_ph['stats']['rms_periods'])
            step = 1
            if n_periods > 200:
                step = max(1, round(n_periods / 20))
                step = max(1, (step // 5) * 5)
            indices = list(range(0, n_periods, step))

            rms_table_header = ["№ периода"] + [f"Фаза {ph['name']}, В" for ph in phases]
            table4 = doc.add_table(rows=1, cols=len(rms_table_header))
            table4.style = 'Table Grid'
            table4.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, hdr in enumerate(rms_table_header):
                cell = table4.rows[0].cells[i]
                cell.text = hdr
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0,0,0)

            for idx in indices:
                row_cells = table4.add_row().cells
                row_cells[0].text = str(idx+1)
                for j, ph in enumerate(phases):
                    if idx < len(ph['stats']['rms_periods']):
                        row_cells[j+1].text = f"{ph['stats']['rms_periods'][idx]:.2f}"
                    else:
                        row_cells[j+1].text = ""
                for cell in row_cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0,0,0)
            self._format_table(table4)

        # ---------- Графики (фото) ----------
        self._add_heading_word(doc, "Графики", level=2)
        plot_files = [
            "Сравнение_отн_амплитуд_гармоник.png",
            "RMS_периоды_фаза_А.png", "RMS_периоды_фаза_В.png", "RMS_периоды_фаза_С.png",
            "Гармоники_1-40_ГОСТ_фаза_А.png", "Гармоники_1-40_ГОСТ_фаза_В.png", "Гармоники_1-40_ГОСТ_фаза_С.png",
            "Гармоники_2-40_ГОСТ_фаза_А.png", "Гармоники_2-40_ГОСТ_фаза_В.png", "Гармоники_2-40_ГОСТ_фаза_С.png",
            "Амплитудный_спектр_фаза_А.png", "Амплитудный_спектр_фаза_В.png", "Амплитудный_спектр_фаза_С.png",
            "Сигнал_фаза_А.png", "Сигнал_фаза_В.png", "Сигнал_фаза_С.png",
            "Общий_анализ_фаза_А.png", "Общий_анализ_фаза_В.png", "Общий_анализ_фаза_С.png"
        ]
        for fname in plot_files:
            fpath = os.path.join(self.save_folder.get(), fname)
            if os.path.exists(fpath):
                doc.add_picture(fpath, width=Inches(5.5))
                self._add_paragraph_word(doc, os.path.basename(fname))

        self._save_word(doc, "Отчет_напряжение_расширенный.docx")

    def generate_extended_current_word(self, data):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as docx_qn
        channels = data['channels']
        long_analysis = data['long_analysis']
        tInc = data['tInc']

        for ch in channels:
            doc = self._setup_word_document()
            # Убираем режим совместимости
            settings = doc.settings.element
            for child in settings:
                if child.tag == docx_qn('w:compat'):
                    settings.remove(child)
                    break

            self._add_heading_word(doc, "Отчёт по току")
            self._add_paragraph_word(doc, "")

            # Общие сведения
            self._add_heading_word(doc, "Общие сведения", level=2)
            self._add_paragraph_word(doc, f"Дата и время анализа: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
            self._add_paragraph_word(doc, f"Основная частота: {ch['f0']:.3f} Гц")
            self._add_paragraph_word(doc, f"Частота дискретизации: {1.0/ch['tInc']:.2f} Гц")
            self._add_paragraph_word(doc, f"Количество отсчётов: {ch['stats']['N_samples']}")
            duration = ch['stats']['N_samples'] * ch['tInc']
            self._add_paragraph_word(doc, f"Длительность сигнала: {duration:.3f} с")
            self._add_paragraph_word(doc, f"Коэффициент масштабирования тока: {self.settings['scale_current']:.3f}")

            # Параметры измерений
            self._add_heading_word(doc, "Параметры измерений", level=2)
            table1 = doc.add_table(rows=1, cols=2)
            table1.style = 'Table Grid'
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            hdr_cells[0].text = "Параметр"
            hdr_cells[1].text = "Значение"
            for cell in hdr_cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

            s = ch['stats']
            rows_data = [
                ("Действующее значение тока Irms", f"{s['Irms']:.3f} А"),
                ("Амплитуда основной гармоники", f"{s['I1_peak']:.3f} А (пик) / {s['I1_rms']:.3f} А (RMS)"),
                ("THDI", f"{s['THDi']:.2f} %"),
                ("Максимальный мгновенный ток", f"{s['Imax_inst']:.3f} А"),
                ("Минимальный мгновенный ток", f"{s['Imin_inst']:.3f} А"),
                ("Пиковый ток", f"{s['Ipeak']:.3f} А"),
                ("Коэффициент амплитуды тока (peak/RMS)", f"{s['crest_factor']:.3f}"),
                ("Анализ по периодам (целых периодов: {})".format(len(s['rms_periods'])), str(len(s['rms_periods']))),
                ("Макс. RMS за период", f"{s['max_rms_period']:.3f} А"),
                ("Мин. RMS за период", f"{s['min_rms_period']:.3f} А"),
                ("Среднее RMS за период", f"{s['avg_rms_period']:.3f} А")
            ]
            for label, value in rows_data:
                row = table1.add_row().cells
                row[0].text = label
                row[1].text = value
                for cell in row:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 0, 0)
            self._format_table(table1)

            # Таблица 2. Гармоники тока
            self._add_heading_word(doc, "Таблица 2. Гармоники тока (1..40)", level=2)
            table2 = doc.add_table(rows=1, cols=5)
            table2.style = 'Table Grid'
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers2 = ["№", "Частота, Гц", "Амплитуда, А", "Отн. ампл., %", "Фаза, °"]
            for i, h in enumerate(headers2):
                cell = table2.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

            I1_rms = s['I1_rms']
            for h in range(1, 41):
                freq = h * ch['f0']
                amp_peak, comp = ch['harmonics'].get(h, (0, 0))
                amp_rms = amp_peak / np.sqrt(2)
                rel = (amp_rms / I1_rms * 100) if I1_rms != 0 else 0
                phase = np.angle(comp, deg=True) if comp != 0 else 0
                row = table2.add_row().cells
                row[0].text = str(h)
                row[1].text = f"{freq:.2f}"
                row[2].text = f"{amp_rms:.4f}"
                row[3].text = f"{rel:.2f}"
                row[4].text = f"{phase:.1f}"
                for cell in row:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 0, 0)
            self._format_table(table2)

            # Вставка графиков
            self._add_heading_word(doc, "Графики", level=2)
            plot_files = [
                f"RMS_тока_периоды_{ch['name']}.png",
                f"Сигнал_тока_{ch['name']}.png",
                f"Гармоники_тока_1-40_{ch['name']}.png",
                f"Гармоники_тока_2-40_{ch['name']}.png",
                f"THDi_диаграмма_{ch['name']}.png",
                f"Спектр_тока_{ch['name']}.png",
                f"Общий_анализ_тока_{ch['name']}.png"
            ]
            for fname in plot_files:
                fpath = os.path.join(self.save_folder.get(), fname)
                if os.path.exists(fpath):
                    doc.add_picture(fpath, width=Inches(5.5))
                    self._add_paragraph_word(doc, os.path.basename(fname))

            self._save_word(doc, f"Отчет_ток_{ch['name']}.docx")






    # ========== EXCEL отчёты ==========
    def generate_extended_voltage_excel(self, data):
        phases = data['phases']
        long_analysis = data['long_analysis']
        tInc = data['tInc']
        ch_names = data['ch_names']
        assignments = data['assignments']

        wb = openpyxl.Workbook()
        font_title = Font(name='Times New Roman', size=14, bold=True)
        font_main = Font(name='Times New Roman', size=14)
        align_center = Alignment(horizontal='center', vertical='center')

        # Лист "Общие сведения"
        ws_info = wb.active
        ws_info.title = "Общие сведения"
        info_data = [
            ["Параметр", "Значение"],
            ["Дата и время", datetime.now().strftime('%d.%m.%Y %H:%M:%S')],
            ["Формат", "Rigol DHO814"],
            ["Частота дискретизации, Гц", 1.0/tInc],
            ["Количество отсчётов", len(phases[0]['signal'])],
            ["Масштаб напряжения", self.settings['scale_voltage']],
            ["Масштаб тока", self.settings['scale_current']],
            ["Канал(ы) напряжения", ", ".join([ch_names[i] for i, a in enumerate(assignments) if a == "Напряжение"])],
            ["Количество фаз", len(phases)]
        ]
        self._write_sheet_data(ws_info, info_data, font_main, align_center)

        # Лист "Параметры напряжения"
        ws_volt = wb.create_sheet("Параметры напряжения")
        headers = ["Параметр"] + [f"Фаза {ph['name']}" for ph in phases] + (["Ср."] if len(phases)>1 else [])
        self._write_excel_row(ws_volt, 1, headers, font_title, align_center)
        row = 2
        s_list = [ph['stats'] for ph in phases]
        def add_row_ws(name, values, unit=""):
            nonlocal row
            r = [name] + [f"{v:.2f}{unit}" for v in values]
            if len(phases) > 1: r.append(f"{np.mean(values):.2f}{unit}")
            self._write_excel_row(ws_volt, row, r, font_main, align_center)
            row += 1
        add_row_ws("Urms, В", [s['Urms'] for s in s_list])
        add_row_ws("Отклонение, %", [s['deviation_percent'] for s in s_list])
        add_row_ws("U макс. мгн., В", [s['Umax_inst'] for s in s_list])
        add_row_ws("U мин. мгн., В", [s['Umin_inst'] for s in s_list])
        add_row_ws("U пик., В", [s['Upeak'] for s in s_list])
        add_row_ws("Коэф. ампл.", [s['crest_factor'] for s in s_list])
        add_row_ws("Макс. RMS пер.", [s['max_rms_period'] for s in s_list])
        add_row_ws("Мин. RMS пер.", [s['min_rms_period'] for s in s_list])
        add_row_ws("Ср. RMS пер.", [s['avg_rms_period'] for s in s_list])

        # Листы по фазам (гармоники)
        for ph in phases:
            ws_phase = wb.create_sheet(f"Фаза {ph['name']}")
            self._write_excel_row(ws_phase, 1, ["Гармоника", "Частота, Гц", "Амплитуда, В", "Отн. ампл., %", "Предел ГОСТ, %"], font_title, align_center)
            row = 2
            U1 = ph['stats']['Urms']
            for h in range(1, 41):
                freq = h * ph['f0']
                amp = ph['rms_harm'].get(h, 0)
                rel = (amp / U1 * 100) if U1 else 0
                limit = GOST_HARM_ALL.get(h, 0)
                self._write_excel_row(ws_phase, row, [h, f"{freq:.1f}", f"{amp:.2f}", f"{rel:.2f}", f"{limit:.2f}"], font_main, align_center)
                row += 1

        # Лист RMS по периодам (если длительный анализ)
        if long_analysis:
            ws_rms = wb.create_sheet("RMS по периодам")
            headers_rms = ["Период"] + [f"Фаза {ph['name']}, В" for ph in phases]
            self._write_excel_row(ws_rms, 1, headers_rms, font_title, align_center)
            ref_ph = phases[0]
            n_periods = len(ref_ph['stats']['rms_periods'])
            step = 1
            if n_periods > 200:
                step = max(1, round(n_periods / 20))
                step = max(1, (step // 5) * 5)
            row = 2
            for idx in range(0, n_periods, step):
                r = [idx+1]
                for ph in phases:
                    if idx < len(ph['stats']['rms_periods']):
                        r.append(f"{ph['stats']['rms_periods'][idx]:.2f}")
                    else:
                        r.append("")
                self._write_excel_row(ws_rms, row, r, font_main, align_center)
                row += 1

        # Лист "График"
        ws_chart = wb.create_sheet("График")
        # Данные с гармоники 2
        chart_headers = ["Гармоника", "Отн. ампл. ср., %", "Предел ГОСТ, %"]
        self._write_excel_row(ws_chart, 1, chart_headers, font_title, align_center)
        row_chart = 2
        for h in range(2, 40):   # гармоники 2–39 (или до 40)
            rels = []
            for ph in phases:
                U1 = ph['stats']['Urms']
                amp = ph['rms_harm'].get(h, 0)
                rel = (amp / U1 * 100) if U1 else 0
                rels.append(rel)
            avg_rel = np.mean(rels)
            limit = GOST_HARM_ALL.get(h, 0)
            self._write_excel_row(ws_chart, row_chart, [h, avg_rel, limit], font_main, align_center)
            row_chart += 1

        # Создаём комбинированную диаграмму: столбцы для Отн. ампл. ср., и точечная с линией для предела
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "Гармонический состав и пределы ГОСТ (без 1-й гармоники)"
        chart.y_axis.title = "%"
        chart.x_axis.title = "Номер гармоники"
        # Данные столбцов – первый ряд (B2:Bxx)
        data_ref = Reference(ws_chart, min_col=2, min_row=1, max_row=row_chart-1)
        cats_ref = Reference(ws_chart, min_col=1, min_row=2, max_row=row_chart-1)
        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cats_ref)
        chart.series[0].graphicalProperties.solidFill = "4472C4"

        # Добавляем второй ряд – предел ГОСТ как линейный график на вторичной оси? Лучше на той же оси.
        # Создаём объект Series для линейного графика
        from openpyxl.chart.series import DataPoint
        from openpyxl.chart import Reference as ChartRef
        # В openpyxl можно добавить ряд данных другого типа на ту же диаграмму, но проще создать вторую диаграмму и объединить.
        # Для простоты сделаем два отдельных графика, но по условию нужно на одном.
        # Используем overlay: создаём LineChart на той же системе координат.
        from openpyxl.chart import LineChart, Series
        line_chart = LineChart()
        line_chart.add_data(Reference(ws_chart, min_col=3, min_row=1, max_row=row_chart-1), titles_from_data=True)
        line_chart.series[0].graphicalProperties.line.solidFill = "ED7D31"
        line_chart.series[0].graphicalProperties.line.width = 25000  # толщина
        line_chart.y_axis.axId = 200
        # Устанавливаем общий диапазон оси X
        line_chart.set_categories(cats_ref)
        # Объединяем
        chart += line_chart
        # Настройка осей
        chart.y_axis.crosses = "min"
        ws_chart.add_chart(chart, "E2")

        # Сохранение
        path = os.path.join(self.save_folder.get(), "Отчет_напряжение_расширенный.xlsx")
        wb.save(path)
        self.log(f"Excel отчёт сохранён: {path}")

    def _write_sheet_data(self, ws, data, font, alignment):
        for r, row in enumerate(data, 1):
            for c, value in enumerate(row, 1):
                cell = ws.cell(row=r, column=c, value=value)
                cell.font = font
                cell.alignment = alignment
                if isinstance(value, float):
                    cell.number_format = '0.00'

    def _write_excel_row(self, ws, row, values, font, alignment):
        for c, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=c, value=val)
            cell.font = font
            cell.alignment = alignment
            if isinstance(val, float):
                cell.number_format = '0.00'

    def generate_extended_current_excel(self, data):
        channels = data['channels']
        for ch in channels:
            wb = openpyxl.Workbook()
            font_main = Font(name='Times New Roman', size=14)
            align_center = Alignment(horizontal='center', vertical='center')

            # Общие сведения
            ws_info = wb.active
            ws_info.title = "Общие сведения"
            info = [
                ["Параметр", "Значение"],
                ["Дата и время анализа", datetime.now().strftime('%d.%m.%Y %H:%M:%S')],
                ["Основная частота, Гц", ch['f0']],
                ["Частота дискретизации, Гц", 1.0 / ch['tInc']],
                ["Количество отсчётов", ch['stats']['N_samples']],
                ["Длительность сигнала, с", ch['stats']['N_samples'] * ch['tInc']],
                ["Коэффициент масштабирования тока", self.settings['scale_current']]
            ]
            self._write_sheet_data(ws_info, info, font_main, align_center)

            # Параметры измерений
            ws_param = wb.create_sheet("Параметры измерений")
            s = ch['stats']
            param_data = [
                ["Параметр", "Значение"],
                ["Действующее значение тока Irms, А", s['Irms']],
                ["Амплитуда основной гармоники (пик), А", s['I1_peak']],
                ["THDI, %", s['THDi']],
                ["Максимальный мгновенный ток, А", s['Imax_inst']],
                ["Минимальный мгновенный ток, А", s['Imin_inst']],
                ["Пиковый ток, А", s['Ipeak']],
                ["Коэффициент амплитуды тока", s['crest_factor']],
                ["Целых периодов", len(s['rms_periods'])],
                ["Макс. RMS за период, А", s['max_rms_period']],
                ["Мин. RMS за период, А", s['min_rms_period']],
                ["Среднее RMS за период, А", s['avg_rms_period']]
            ]
            self._write_sheet_data(ws_param, param_data, font_main, align_center)

            # Гармоники тока
            ws_harm = wb.create_sheet("Гармоники тока")
            headers = ["№", "Частота, Гц", "Амплитуда, А", "Отн. ампл., %", "Фаза, °"]
            self._write_excel_row(ws_harm, 1, headers, font_main, align_center)
            I1_rms = s['I1_rms']
            for h in range(1, 41):
                freq = h * ch['f0']
                amp_peak, comp = ch['harmonics'].get(h, (0, 0))
                amp_rms = amp_peak / np.sqrt(2)
                rel = (amp_rms / I1_rms * 100) if I1_rms else 0
                phase = np.angle(comp, deg=True)
                self._write_excel_row(ws_harm, h + 1, [h, freq, amp_rms, rel, phase], font_main, align_center)

            # Автоширина столбцов
            for ws in [ws_info, ws_param, ws_harm]:
                for col in ws.columns:
                    max_len = 0
                    col_letter = col[0].column_letter
                    for cell in col:
                        try:
                            if cell.value:
                                max_len = max(max_len, len(str(cell.value)))
                        except:
                            pass
                    ws.column_dimensions[col_letter].width = min(max_len + 2, 35)

            path = os.path.join(self.save_folder.get(), f"Отчет_ток_{ch['name']}.xlsx")
            wb.save(path)
            self.log(f"Excel отчёт по току сохранён: {path}")




    # ========== ГРАФИКИ (фото) ==========
    def generate_extended_voltage_plots(self, data):
        phases = data['phases']
        long_analysis = data['long_analysis']
        # Для каждой фазы строим множество графиков
        phase_names = [ph['name'] for ph in phases]
        # 1. Сравнение относительных амплитуд гармоник по фазам
        self._plot_relative_harmonics_comparison(phases)
        # 2. RMS по периодам для каждой фазы 
        for ph in phases:
            self._plot_rms_periods(ph)
        # 3. Гармоники 1-40 с пределами ГОСТ для каждой фазы
        for ph in phases:
            self._plot_harmonics_with_gost(ph, start_harm=1, end_harm=40)
        # 4. Гармоники 2-40 с пределами ГОСТ для каждой фазы
        for ph in phases:
            self._plot_harmonics_with_gost(ph, start_harm=2, end_harm=40)
        # 5. Амплитудный спектр (без чисел)
        for ph in phases:
            self._plot_spectrum(ph)
        # 6. Сигнал (временная диаграмма)
        for ph in phases:
            self._plot_signal(ph)
        # 7. Общий анализ гармоник фазы (5 subplots)
        for ph in phases:
            self._plot_combined_analysis(ph)

    def _plot_relative_harmonics_comparison(self, phases):
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        harm_range = range(2, 41)   # без первой гармоники
        x = np.arange(len(harm_range))
        width = 0.22
        max_rel = 0
        for i, ph in enumerate(phases):
            U1 = ph['stats']['Urms']
            rels = [ (ph['rms_harm'].get(h,0)/U1*100) if U1 else 0 for h in harm_range]
            offset = (i - 1) * width
            bars = ax.bar(x + offset, rels, width, label=f"Фаза {ph['name']}")
            for bar, rel in zip(bars, rels):
                if rel > max_rel:
                    max_rel = rel
                                    
        ax.set_xticks(x)
        ax.set_xticklabels([str(h) for h in harm_range])
        ax.set_xlabel("Номер гармоники")
        ax.set_ylabel("Относительная амплитуда, %")
        ax.set_title("Сравнение относительных амплитуд гармоник по фазам (без 1-й)")
        ax.legend()
        ax.grid(True)
        # Немного увеличим верхнюю границу, чтобы подписи не обрезались
        ax.set_ylim(0, max_rel * 1.15 if max_rel > 0 else 10)
        self._save_figure(fig, "Сравнение_отн_амплитуд_гармоник.png")

    def _plot_rms_periods(self, ph):
        long_analysis = self.long_analysis_var.get()
        times = ph['stats']['times_period']
        rms_vals = ph['stats']['rms_periods']
        if len(times) < 2:
            return
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        ax.plot(times, rms_vals, 'b-o', markersize=4)
        nom = self.settings['nominal_voltage']
        ax.axhline(nom, color='r', linestyle='--', label='Номинал')
        ax.axhline(nom*0.97, color='g', linestyle='--', label='−3%')
        ax.axhline(nom*1.03, color='g', linestyle='--', label='+3%')
        ax.set_ylim(nom*0.9, nom*1.1)

        # Подписи для каждого периода (только если периодов мало)
        if not long_analysis and len(rms_vals) <= 20:
            for t, val in zip(times, rms_vals):
                ax.text(t, val + 0.3, f"{val:.1f}", ha='center', fontsize=6)

        # Статистика (min, max, среднее) – текст в левом верхнем углу, без лишних линий
        min_v, max_v, mean_v = np.min(rms_vals), np.max(rms_vals), np.mean(rms_vals)
        stats_text = f"min: {min_v:.1f} В\nmax: {max_v:.1f} В\nmean: {mean_v:.1f} В"
        props = dict(boxstyle='round', facecolor='wheat', alpha=0.5)
        ax.text(0.02, 0.98, stats_text, transform=ax.transAxes, fontsize=9,
                verticalalignment='top', bbox=props)

        ax.set_xlabel("Время, с")
        ax.set_ylabel("RMS, В")
        ax.set_title(f"RMS по периодам (±3%) - Фаза {ph['name']}")
        ax.grid(True)
        ax.legend()
        self._save_figure(fig, f"RMS_периоды_фаза_{ph['name']}.png")

    def _plot_harmonics_with_gost(self, ph, start_harm=1, end_harm=40):
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        U1 = ph['stats']['Urms']
        h_list = list(range(start_harm, end_harm+1))
        rels = [ (ph['rms_harm'].get(h,0)/U1*100) if U1 else 0 for h in h_list]
        limits = [GOST_HARM_ALL.get(h, 0) for h in h_list]

        # Строим столбцы предела (полупрозрачные) фоном
        limit_bars = ax.bar(h_list, limits, color='red', alpha=0.2, label='Предел ГОСТ')
        # Столбцы относительной амплитуды: красный, если превышен предел
        bar_colors = ['red' if r > l and l > 0 else 'steelblue' for r, l in zip(rels, limits)]
        bars = ax.bar(h_list, rels, color=bar_colors, label='Отн. ампл., %')

        # Подписи значений горизонтально
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.1, f'{height:.2f}',
                        ha='center', va='bottom', fontsize=7)
        ax.set_xlabel("Номер гармоники")
        ax.set_ylabel("%")
        ax.set_title(f"Гармоники {start_harm}-{end_harm} с пределами ГОСТ - Фаза {ph['name']}")
        ax.legend()
        ax.grid(True)
        self._save_figure(fig, f"Гармоники_{start_harm}-{end_harm}_ГОСТ_фаза_{ph['name']}.png") 
        
    def _plot_spectrum(self, ph):
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        signal = ph['signal']
        tInc = ph['tInc']
        N = len(signal)
        window = np.hanning(N)
        y = signal * window
        Y = rfft(y)
        freqs = rfftfreq(N, tInc)
        ax.stem(freqs, np.abs(Y), markerfmt=' ', basefmt=' ')
        ax.set_xlim(0, 2000)  # до 2 кГц
        ax.set_xlabel("Частота, Гц")
        ax.set_ylabel("Амплитуда")
        ax.set_title(f"Амплитудный спектр - Фаза {ph['name']}")
        ax.grid(True)
        # без чисел над столбиками
        self._save_figure(fig, f"Амплитудный_спектр_фаза_{ph['name']}.png")

    def _plot_signal(self, ph):
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        t = np.arange(len(ph['signal'])) * ph['tInc']
        ax.plot(t, ph['signal'], 'b')
        ax.set_xlabel("Время, с")
        ax.set_ylabel("Напряжение, В")
        ax.set_title(f"Сигнал - Фаза {ph['name']}")
        ax.grid(True)
        self._save_figure(fig, f"Сигнал_фаза_{ph['name']}.png")

    def _plot_combined_analysis(self, ph):
        fig = Figure(figsize=(20, 12), dpi=300)
        axs = fig.subplots(3, 2)
        # Если есть лишний subplot, удалим
        if len(axs.shape) > 1 and axs.shape[1] == 2:
            try:
                fig.delaxes(axs[2,1])
            except:
                pass

        # 1. RMS по периодам (всегда строим, если возможно) или просто сигнал
        long_analysis = self.long_analysis_var.get()
        times_period, rms_vals = ph['stats']['times_period'], ph['stats']['rms_periods']
        ax = axs[0,0]
        if len(times_period) >= 2:
            ax.plot(times_period, rms_vals, 'b-o', markersize=4)
            # Горизонтальные линии номинала и ±3%
            nom = self.settings['nominal_voltage']
            ax.axhline(nom, color='r', linestyle='--', label='Номинал')
            ax.axhline(nom*0.97, color='g', linestyle='--', label='-3%')
            ax.axhline(nom*1.03, color='g', linestyle='--', label='+3%')
            ax.set_ylim(nom*0.9, nom*1.1)  # развертка ±10% от номинала
            # Подписи значений
            if not long_analysis and len(rms_vals) <= 20:
                for i, val in enumerate(rms_vals):
                    ax.text(times_period[i], val + 0.5, f"{val:.1f}", ha='center', fontsize=6)
            else:
                # только мин, макс, среднее
                min_val, max_val, mean_val = np.min(rms_vals), np.max(rms_vals), np.mean(rms_vals)
                ax.axhline(min_val, color='m', linestyle='--')
                ax.axhline(max_val, color='m', linestyle='--')
                ax.axhline(mean_val, color='k', linestyle='--')
                ax.text(times_period[0], min_val, f'{min_val:.1f}', va='bottom', fontsize=6)
                ax.text(times_period[0], max_val, f'{max_val:.1f}', va='top', fontsize=6)
                ax.text(times_period[-1], mean_val, f'{mean_val:.1f}', va='bottom', fontsize=6)
            ax.set_title("RMS по периодам")
            ax.grid(True)
        else:
            # fallback на сигнал
            t = np.arange(len(ph['signal'])) * ph['tInc']
            ax.plot(t, ph['signal'])
            ax.set_title("Сигнал (недостаточно периодов)")

        # 2. Гармоники 1-40 с ГОСТ
        ax = axs[0,1]
        self._draw_harmonics_gost_on_ax(ax, ph, 1, 40)

        # 3. Гармоники 2-40 с ГОСТ
        ax = axs[1,0]
        self._draw_harmonics_gost_on_ax(ax, ph, 2, 40)

        # 4. Амплитудный спектр
        ax = axs[1,1]
        signal = ph['signal']; tInc = ph['tInc']; N = len(signal)
        window = np.hanning(N)
        Y = rfft(signal * window)
        freqs = rfftfreq(N, tInc)
        ax.stem(freqs, np.abs(Y), markerfmt=' ', basefmt=' ')
        ax.set_xlim(0, 2000)
        ax.set_title("Амплитудный спектр")

        # 5. Сигнал
        ax = axs[2,0]
        t = np.arange(N) * tInc
        ax.plot(t, signal, 'b')
        ax.set_title("Сигнал")

        fig.tight_layout()
        self._save_figure(fig, f"Общий_анализ_фаза_{ph['name']}.png")

    def _draw_harmonics_gost_on_ax(self, ax, ph, start, end):
        U1 = ph['stats']['Urms']
        h_list = list(range(start, end+1))
        rels = [ (ph['rms_harm'].get(h,0)/U1*100) if U1 else 0 for h in h_list]
        limits = [GOST_HARM_ALL.get(h, 0) for h in h_list]
        ax.bar(h_list, limits, color='red', alpha=0.2, label='Предел ГОСТ')
        bar_colors = ['red' if r > l and l > 0 else 'steelblue' for r, l in zip(rels, limits)]
        ax.bar(h_list, rels, color=bar_colors, label='Отн. ампл., %')
        ax.set_title(f"Гармоники {start}-{end}")
        ax.legend(fontsize='small')
        ax.grid(True)

    def _save_figure(self, fig, filename):
        path = os.path.join(self.save_folder.get(), filename)
        fig.savefig(path, bbox_inches='tight')
        self.log(f"График сохранён: {path}")

    def generate_extended_current_plots(self, data):
        channels = data['channels']
        long_analysis = data['long_analysis']
        for ch in channels:
            self._plot_current_rms_periods(ch, long_analysis)
            self._plot_current_signal(ch)
            self._plot_current_harmonics(ch, 1, 40)
            self._plot_current_harmonics(ch, 2, 40)
            self._plot_current_thdi_pie(ch)
            self._plot_current_spectrum(ch)
            self._plot_current_combined(ch)

    def _plot_current_rms_periods(self, ch, long_analysis):
        times = ch['stats']['times_period']
        rms_vals = ch['stats']['rms_periods']
        if len(times) < 2:
            return
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        ax.plot(times, rms_vals, 'b-o', markersize=4)

        # Автоматический запас для подписей
        y_min, y_max = np.min(rms_vals), np.max(rms_vals)
        y_range = y_max - y_min if y_max > y_min else 1
        ax.set_ylim(y_min - 0.1 * y_range, y_max + 0.15 * y_range)

        if not long_analysis and len(rms_vals) <= 20:
            for t, val in zip(times, rms_vals):
                ax.text(t, val + 0.03 * y_range, f"{val:.2f}", ha='center', fontsize=6)
        else:
            min_v, max_v, mean_v = y_min, y_max, np.mean(rms_vals)
            stats_text = f"min: {min_v:.2f} А\nmax: {max_v:.2f} А\nmean: {mean_v:.2f} А"
            props = dict(boxstyle='round', facecolor='wheat', alpha=0.5)
            ax.text(0.02, 0.98, stats_text, transform=ax.transAxes, fontsize=9,
                    verticalalignment='top', bbox=props)

        ax.set_xlabel("Время, с")
        ax.set_ylabel("RMS, А")
        ax.set_title(f"RMS тока по периодам - {ch['name']}")
        ax.grid(True)
        self._save_figure(fig, f"RMS_тока_периоды_{ch['name']}.png")

    def _plot_current_signal(self, ch):
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        t = np.arange(len(ch['signal'])) * ch['tInc']
        ax.plot(t, ch['signal'], 'b')
        ax.set_xlabel("Время, с")
        ax.set_ylabel("Ток, А")
        ax.set_title(f"Исходный сигнал тока (DC удалён) - {ch['name']}")
        ax.grid(True)
        self._save_figure(fig, f"Сигнал_тока_{ch['name']}.png")

    def _plot_current_harmonics(self, ch, start, end):
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        h_list = list(range(start, end + 1))
        I1_rms = ch['stats']['I1_rms']
        rels = []
        for h in h_list:
            amp_peak, _ = ch['harmonics'].get(h, (0, 0))
            amp_rms = amp_peak / np.sqrt(2)
            rel = (amp_rms / I1_rms * 100) if I1_rms else 0
            rels.append(rel)
        bars = ax.bar(h_list, rels, color='steelblue')
        ax.set_xlabel("Номер гармоники")
        ax.set_ylabel("Относительная амплитуда, %")
        ax.set_title(f"Гармоники тока {start}-{end} - {ch['name']}")
        ax.grid(True)
        for bar in bars:
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width() / 2., h, f'{h:.2f}', ha='center', va='bottom', fontsize=7)
        self._save_figure(fig, f"Гармоники_тока_{start}-{end}_{ch['name']}.png")

    def _plot_current_thdi_pie(self, ch):
        fig = Figure(figsize=(12, 10), dpi=300)
        ax = fig.add_subplot(111)
        squares = []
        labels = []
        for h in range(2, 41):
            amp_peak, _ = ch['harmonics'].get(h, (0, 0))
            amp_rms = amp_peak / np.sqrt(2)
            if amp_rms > 0:
                squares.append(amp_rms ** 2)
                labels.append(str(h))
        total_sq = sum(squares)
        if total_sq == 0:
            ax.text(0.5, 0.5, 'THDi = 0%', transform=ax.transAxes, ha='center')
            ax.axis('off')
        else:
            percentages = [s / total_sq * 100 for s in squares]
            wedges, texts, autotexts = ax.pie(
                percentages, labels=None, autopct='%1.1f%%',
                startangle=90, pctdistance=0.85,
                textprops={'fontsize': 8}
            )
            # Выносим номера гармоник за пределы диаграммы
            for i, (wedge, label) in enumerate(zip(wedges, labels)):
                ang = (wedge.theta2 - wedge.theta1) / 2. + wedge.theta1
                x = 1.2 * np.cos(np.deg2rad(ang))
                y = 1.2 * np.sin(np.deg2rad(ang))
                ax.text(x, y, label, ha='center', va='center', fontsize=9)
            ax.set_title(f"Вклад гармоник в THDi - {ch['name']}")
        self._save_figure(fig, f"THDi_диаграмма_{ch['name']}.png")

    def _plot_current_spectrum(self, ch):
        fig = Figure(figsize=(14, 8), dpi=300)
        ax = fig.add_subplot(111)
        sig = ch['signal']
        N = len(sig)
        window = np.hanning(N)
        y = sig * window
        Y = rfft(y)
        freqs = rfftfreq(N, ch['tInc'])
        ax.stem(freqs, np.abs(Y), markerfmt=' ', basefmt=' ')
        ax.set_xlim(0, 2000)
        ax.set_xlabel("Частота, Гц")
        ax.set_ylabel("Амплитуда")
        ax.set_title(f"Амплитудный спектр тока - {ch['name']}")
        ax.grid(True)
        self._save_figure(fig, f"Спектр_тока_{ch['name']}.png")

    def _plot_current_combined(self, ch):
        fig = Figure(figsize=(20, 12), dpi=300)
        axs = fig.subplots(2, 2)

        # Сигнал
        ax = axs[0, 0]
        t = np.arange(len(ch['signal'])) * ch['tInc']
        ax.plot(t, ch['signal'], 'b')
        ax.set_title("Сигнал тока (DC удалён)")
        ax.grid(True)

        # RMS по периодам
        ax = axs[0, 1]
        times = ch['stats']['times_period']
        rms_vals = ch['stats']['rms_periods']
        if len(times) >= 2:
            ax.plot(times, rms_vals, 'b-o', markersize=3)
            y_min, y_max = np.min(rms_vals), np.max(rms_vals)
            y_range = y_max - y_min if y_max > y_min else 1
            ax.set_ylim(y_min - 0.1 * y_range, y_max + 0.15 * y_range)
            ax.set_title("RMS по периодам")
            ax.grid(True)
        else:
            ax.text(0.5, 0.5, "Нет данных", transform=ax.transAxes, ha='center')

        # Круговая диаграмма THDi
        ax = axs[1, 0]
        squares = []
        labels = []
        for h in range(2, 41):
            amp_peak, _ = ch['harmonics'].get(h, (0, 0))
            amp_rms = amp_peak / np.sqrt(2)
            if amp_rms > 0:
                squares.append(amp_rms ** 2)
                labels.append(str(h))
        total_sq = sum(squares)
        if total_sq > 0:
            percentages = [s / total_sq * 100 for s in squares]
            wedges, texts, autotexts = ax.pie(
                percentages, labels=None, autopct='%1.1f%%',
                startangle=90, pctdistance=0.85,
                textprops={'fontsize': 6}
            )
            for i, (wedge, label) in enumerate(zip(wedges, labels)):
                ang = (wedge.theta2 - wedge.theta1) / 2. + wedge.theta1
                x = 1.25 * np.cos(np.deg2rad(ang))
                y = 1.25 * np.sin(np.deg2rad(ang))
                ax.text(x, y, label, ha='center', va='center', fontsize=7)
            ax.set_title("Вклад гармоник в THDi")
        else:
            ax.text(0.5, 0.5, "THDi = 0%", transform=ax.transAxes, ha='center')
            ax.axis('off')

        # Спектр
        ax = axs[1, 1]
        sig = ch['signal']
        N = len(sig)
        window = np.hanning(N)
        Y = rfft(sig * window)
        freqs = rfftfreq(N, ch['tInc'])
        ax.stem(freqs, np.abs(Y), markerfmt=' ', basefmt=' ')
        ax.set_xlim(0, 2000)
        ax.set_title("Амплитудный спектр")
        ax.grid(True)

        fig.tight_layout()
        self._save_figure(fig, f"Общий_анализ_тока_{ch['name']}.png")




    # -------------------- Обработчики других режимов (оставлены для совместимости) --------------------
    def display_current_results(self, results):
        report = "=== Результаты анализа тока ===\n"
        for ch, res in results.items():
            report += f"\nКанал {ch}:\n  Частота сети: {res['f0']:.2f} Гц\n  I1 RMS: {res['fundamental_rms']:.3f} А\n  THDi: {res['thd']:.2f}%\n"
        self.log(report)
        if self.have_plots and results:
            first_ch = list(results.values())[0]
            self.plot_harmonics(first_ch['harmonic_rms'], "Ток", "А")
            self.save_plot_image("График_гармоник_тока.png")
        if HAVE_DOCX:
            self.generate_current_report_word(results)
        if HAVE_OPENPYXL:
            self.export_harmonics_to_excel(results, "Данные_гармоник_тока.xlsx", "Ток")

    def display_combined_results(self, res):
        v = res['voltage']; i = res['current']; p = res['power']
        report = "=== Результаты анализа (U + I одной фазы) ===\n"
        report += f"\nНапряжение:\n  f0={v['f0']:.2f} Гц, U1={v['fundamental_rms']:.3f} В, THD={v['thd']:.2f}%\n"
        if v.get('gost_violations'):
            for viol in v['gost_violations']:
                report += f"    Нарушение: {viol}\n"
        report += f"\nТок:\n  I1={i['fundamental_rms']:.3f} А, THDi={i['thd']:.2f}%\n"
        report += f"\nМощности:\n  P={p['P_total']:.2f} Вт, Q={p['Q_total']:.2f} вар, S={p['S_total']:.2f} ВА, PF={p['pf']:.3f}\n"
        self.log(report)
        if self.have_plots:
            self.plot_combined(v['harmonic_rms'], i['harmonic_rms'])
            self.save_plot_image("График_гармоник_U_I.png")
        if HAVE_DOCX:
            self.generate_combined_report_word(res)
        if HAVE_OPENPYXL:
            self.export_combined_excel(v, i, p)

    def display_three_phase_results(self, data):
        report = "=== Результаты трёхфазного анализа ===\n"
        for ph in data['phases']:
            report += f"\nФаза {ph['phase']}:\n"
            report += f"  U1={ph['voltage']['fundamental_rms']:.3f} В, THD={ph['voltage']['thd']:.2f}%\n"
            report += f"  I1={ph['current']['fundamental_rms']:.3f} А, THDi={ph['current']['thd']:.2f}%\n"
            report += f"  P={ph['power']['P_total']:.2f} Вт, Q={ph['power']['Q_total']:.2f} вар, S={ph['power']['S_total']:.2f} ВА, PF={ph['power']['pf']:.3f}\n"
        tp = data['total_power']
        report += f"\nСуммарно:\n  P={tp['P_total']:.2f} Вт, Q={tp['Q_total']:.2f} вар, S={tp['S_total']:.2f} ВА, PF={tp['pf']:.3f}\n"
        report += f"Общий THDv (макс): {data['overall_thd_v']:.2f}%, THDi (макс): {data['overall_thd_i']:.2f}%\n"
        self.log(report)
        if self.have_plots and data['phases']:
            v_h = data['phases'][0]['voltage']['harmonic_rms']
            i_h = data['phases'][0]['current']['harmonic_rms']
            self.plot_combined(v_h, i_h)
            self.save_plot_image("График_трёхфазный.png")
        if HAVE_DOCX:
            self.generate_three_phase_report_word(data)
        if HAVE_OPENPYXL:
            self.export_three_phase_excel(data)

    # ========== Методы для старых отчётов Word (оставлены) ==========
    def _setup_word_document(self):
        doc = Document()
        section = doc.sections[0]
        section.orientation = 1  # landscape
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(0.75)
        section.right_margin = Cm(0.75)
        return doc

    def _add_heading_word(self, doc, text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18) if level==1 else Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.5

    def _add_paragraph_word(self, doc, text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        return p

    def _add_table_word(self, doc, headers, rows, col_widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Cm(width)
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = qn('w:trHeight')
            trPr.set(trHeight, str(int(0.7 * 567)))  # 0.7 cm
        return table

    def _format_table(self, table, header_color='DAE9F7'):
        """Настройка таблицы: повтор заголовка, заливка, высота строк 0.7 см, вертикальное выравнивание."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl_pr = table._tbl.tblPr
        header_el = OxmlElement('w:tblHeader')
        header_el.set(qn('w:val'), 'true')
        tbl_pr.append(header_el)

        for row_idx, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(0.7 * 567)))
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)

            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Вертикальное выравнивание по центру
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                # Горизонтальное выравнивание уже задано через параграфы, но можно и здесь
                # Заливка заголовка
                if row_idx == 0:
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), header_color)
                    tcPr.append(shd)

    def _save_word(self, doc, filename):
        path = os.path.join(self.save_folder.get(), filename)
        doc.save(path)
        self.log(f"Отчёт Word сохранён: {path}")

    # Остальные методы generate_*_word для других режимов оставлены как заглушки, при необходимости можно дополнить
    def generate_current_report_word(self, results):
        doc = self._setup_word_document()
        self._add_heading_word(doc, "Результаты анализа гармоник тока")
        self._add_paragraph_word(doc, f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        for ch, res in results.items():
            self._add_paragraph_word(doc, f"Канал: {ch}", bold=True)
            self._add_paragraph_word(doc, f"f0: {res['f0']:.2f} Гц")
            self._add_paragraph_word(doc, f"I1: {res['fundamental_rms']:.3f} А, THDi: {res['thd']:.2f}%")
            harm = res['harmonic_rms']
            headers = ["№", "I, А", "% от I1"]
            rows = []
            i1 = res['fundamental_rms']
            for h in sorted(harm.keys()):
                a = harm[h]
                rows.append([h, f"{a:.2f}", f"{(a/i1*100):.2f}" if i1 else "0.00"])
            if rows:
                self._add_table_word(doc, headers, rows, col_widths=[1.5,5,5])
        self._save_word(doc, "Отчет_анализ_тока.docx")

    def generate_combined_report_word(self, res):
        doc = self._setup_word_document()
        self._add_heading_word(doc, "Результаты анализа U и I одной фазы")
        v = res['voltage']; i = res['current']; p = res['power']
        self._add_paragraph_word(doc, f"Напряжение: f0={v['f0']:.2f} Гц, U1={v['fundamental_rms']:.3f} В, THD={v['thd']:.2f}%")
        self._add_paragraph_word(doc, f"Ток: I1={i['fundamental_rms']:.3f} А, THDi={i['thd']:.2f}%")
        self._add_paragraph_word(doc, f"Мощности: P={p['P_total']:.2f} Вт, Q={p['Q_total']:.2f} вар, S={p['S_total']:.2f} ВА, PF={p['pf']:.3f}")
        self._save_word(doc, "Отчет_анализ_U_I.docx")

    def generate_three_phase_report_word(self, data):
        doc = self._setup_word_document()
        self._add_heading_word(doc, "Результаты трёхфазного анализа гармоник")
        for ph in data['phases']:
            self._add_paragraph_word(doc, f"Фаза {ph['phase']}")
            v = ph['voltage']; i = ph['current']; p = ph['power']
            self._add_paragraph_word(doc, f"U1={v['fundamental_rms']:.3f} В, THD={v['thd']:.2f}%, I1={i['fundamental_rms']:.3f} А")
            self._add_paragraph_word(doc, f"P={p['P_total']:.2f} Вт, Q={p['Q_total']:.2f} вар, S={p['S_total']:.2f} ВА, PF={p['pf']:.3f}")
        tp = data['total_power']
        self._add_paragraph_word(doc, f"Суммарно: P={tp['P_total']:.2f} Вт, Q={tp['Q_total']:.2f} вар, S={tp['S_total']:.2f} ВА")
        self._save_word(doc, "Отчет_трёхфазный_анализ.docx")

    # Экспорт в Excel для других режимов (оставлены)
    def export_harmonics_to_excel(self, results, filename, sheet_prefix):
        wb, ws, font, align = self._prepare_excel_workbook()
        ws.title = f"{sheet_prefix}_гармоники"
        headers = ["Гармоника", "Амплитуда (RMS)", "% от основной"]
        self._write_excel_header(ws, headers, 1, font, align)
        row = 2
        for ch, res in results.items():
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
            cell = ws.cell(row=row, column=1, value=f"Канал {ch}")
            cell.font = Font(name='Times New Roman', size=14, bold=True)
            cell.alignment = align
            row += 1
            fundamental = res['fundamental_rms']
            for h in sorted(res['harmonic_rms'].keys()):
                val = res['harmonic_rms'][h]
                perc = (val / fundamental * 100) if fundamental != 0 else 0.0
                self._write_excel_data(ws, [[h, val, perc]], row, font, align)
                row += 1
            row += 1
        path = os.path.join(self.save_folder.get(), filename)
        wb.save(path)
        self.log(f"Данные Excel сохранены: {path}")

    def export_combined_excel(self, v_res, i_res, p):
        wb = openpyxl.Workbook()
        font = Font(name='Times New Roman', size=14)
        align = Alignment(horizontal='center', vertical='center')
        ws_v = wb.active; ws_v.title = "Напряжение"
        headers = ["Гармоника", "U, В", "% от U1"]
        self._write_excel_header(ws_v, headers, 1, font, align)
        u1 = v_res['fundamental_rms']
        row = 2
        for h in sorted(v_res['harmonic_rms'].keys()):
            val = v_res['harmonic_rms'][h]
            perc = (val/u1*100) if u1 else 0
            self._write_excel_data(ws_v, [[h, val, perc]], row, font, align)
            row += 1
        ws_i = wb.create_sheet("Ток")
        headers = ["Гармоника", "I, А", "% от I1"]
        self._write_excel_header(ws_i, headers, 1, font, align)
        i1 = i_res['fundamental_rms']
        row = 2
        for h in sorted(i_res['harmonic_rms'].keys()):
            val = i_res['harmonic_rms'][h]
            perc = (val/i1*100) if i1 else 0
            self._write_excel_data(ws_i, [[h, val, perc]], row, font, align)
            row += 1
        ws_p = wb.create_sheet("Мощности")
        self._write_excel_header(ws_p, ["Параметр", "Значение"], 1, font, align)
        self._write_excel_data(ws_p, [["P, Вт", p['P_total']], ["Q, вар", p['Q_total']], ["S, ВА", p['S_total']], ["PF", p['pf']]], 2, font, align)
        path = os.path.join(self.save_folder.get(), "Данные_гармоник_U_I.xlsx")
        wb.save(path)
        self.log(f"Данные Excel сохранены: {path}")

    def export_three_phase_excel(self, data):
        wb = openpyxl.Workbook()
        font = Font(name='Times New Roman', size=14)
        align = Alignment(horizontal='center', vertical='center')
        for ph in data['phases']:
            ws = wb.create_sheet(f"Фаза_{ph['phase']}")
            self._write_excel_header(ws, ["Гармоника", "U, В", "% U1", "I, А", "% I1"], 1, font, align)
            v_h = ph['voltage']['harmonic_rms']; i_h = ph['current']['harmonic_rms']
            u1 = ph['voltage']['fundamental_rms']; i1 = ph['current']['fundamental_rms']
            row = 2
            for h in sorted(set(list(v_h.keys()) + list(i_h.keys()))):
                v_val = v_h.get(h, 0); i_val = i_h.get(h, 0)
                v_perc = (v_val/u1*100) if u1 else 0
                i_perc = (i_val/i1*100) if i1 else 0
                self._write_excel_data(ws, [[h, v_val, v_perc, i_val, i_perc]], row, font, align)
                row += 1
        ws_sum = wb.create_sheet("Суммарно")
        self._write_excel_header(ws_sum, ["Параметр", "Значение"], 1, font, align)
        tp = data['total_power']
        sum_data = [
            ["P, Вт", tp['P_total']], ["Q, вар", tp['Q_total']], ["S, ВА", tp['S_total']],
            ["PF", tp['pf'] if tp['pf'] else 0], ["THDv макс", data['overall_thd_v']], ["THDi макс", data['overall_thd_i']]
        ]
        self._write_excel_data(ws_sum, sum_data, 2, font, align)
        path = os.path.join(self.save_folder.get(), "Данные_трёхфазные.xlsx")
        wb.save(path)
        self.log(f"Данные Excel сохранены: {path}")

    def _prepare_excel_workbook(self):
        wb = openpyxl.Workbook()
        ws = wb.active
        font = Font(name='Times New Roman', size=14, color='000000')
        align = Alignment(horizontal='center', vertical='center')
        return wb, ws, font, align

    def _write_excel_header(self, ws, headers, row, font, align):
        for c, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=c, value=h)
            cell.font = font
            cell.alignment = align

    def _write_excel_data(self, ws, data, start_row, font, align):
        for r, row_data in enumerate(data, start_row):
            for c, value in enumerate(row_data, 1):
                cell = ws.cell(row=r, column=c, value=value)
                cell.font = font
                cell.alignment = align
                if isinstance(value, float):
                    cell.number_format = '0.00'

    def save_plot_image(self, filename):
        if not self.have_plots:
            return
        try:
            current_size = self.fig.get_size_inches()
            current_dpi = self.fig.get_dpi()
            self.fig.set_size_inches(14, 8)
            self.fig.set_dpi(300)
            path = os.path.join(self.save_folder.get(), filename)
            self.fig.savefig(path, bbox_inches='tight')
            self.fig.set_size_inches(current_size)
            self.fig.set_dpi(current_dpi)
            self.canvas.draw()
            self.log(f"График сохранён: {path}")
        except Exception as e:
            self.log(f"Ошибка сохранения графика: {e}")

    # plot_harmonics, plot_combined для preview
    def plot_harmonics(self, harm_rms, ylabel, unit):
        self.ax1.clear()
        self.ax2.clear()
        h_list = sorted(harm_rms.keys())
        vals = [harm_rms[h] for h in h_list]
        bars = self.ax2.bar(h_list, vals, color='steelblue')
        self.ax2.set_xlabel("Номер гармоники")
        self.ax2.set_ylabel(f"{ylabel}, {unit}")
        self.ax2.set_title("Спектр гармоник")
        self.ax2.grid(True, alpha=0.3)
        for bar, val in zip(bars, vals):
            height = bar.get_height()
            self.ax2.text(bar.get_x() + bar.get_width()/2., height,
                          f'{val:.2f}', ha='center', va='bottom', fontsize=8, rotation=90)
        self.canvas.draw()

    def plot_combined(self, v_harm, i_harm):
        self.ax1.clear()
        self.ax2.clear()
        h_list = sorted(v_harm.keys())
        v_vals = [v_harm[h] for h in h_list]
        i_vals = [i_harm[h] for h in h_list]
        bar_v = self.ax2.bar(np.array(h_list)-0.2, v_vals, 0.4, label='U, В', color='coral')
        bar_i = self.ax2.bar(np.array(h_list)+0.2, i_vals, 0.4, label='I, А', color='teal')
        self.ax2.set_xlabel("Номер гармоники")
        self.ax2.set_title("Гармоники U и I")
        self.ax2.legend()
        self.ax2.grid(True, alpha=0.3)
        for bars in [bar_v, bar_i]:
            for bar in bars:
                height = bar.get_height()
                self.ax2.text(bar.get_x() + bar.get_width()/2., height,
                              f'{height:.2f}', ha='center', va='bottom', fontsize=6, rotation=90)
        self.canvas.draw()

if __name__ == "__main__":
    root = tk.Tk()
    app = HarmonicsApp(root)
    root.mainloop()
