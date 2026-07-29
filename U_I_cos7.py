import sys
import os
import re
import time
import numpy as np
from numpy.fft import fft
from scipy.signal import windows
import matplotlib
matplotlib.use('Qt5Agg')
import matplotlib.pyplot as plt
import openpyxl
from openpyxl.drawing.image import Image as XLImage
from openpyxl.chart import ScatterChart, Reference, Series
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QPushButton, QLabel, QLineEdit,
                             QFileDialog, QComboBox, QRadioButton, QGroupBox,
                             QGridLayout, QCheckBox, QTextEdit, QProgressBar,
                             QMessageBox, QDoubleSpinBox, QDialog, QDialogButtonBox)

# ----------------------------------------------------------------------
# Логирование
# ----------------------------------------------------------------------
class Logger:
    def __init__(self, gui_log=None, log_file=None):
        self.gui_log = gui_log
        self.log_file = log_file

    def set_gui(self, gui_text_edit):
        self.gui_log = gui_text_edit

    def set_file(self, filepath):
        self.log_file = filepath
        if filepath:
            dirname = os.path.dirname(filepath)
            if dirname and not os.path.exists(dirname):
                try:
                    os.makedirs(dirname, exist_ok=True)
                except Exception as e:
                    self.log_file = None
                    self.log(f"Не удалось создать папку для лога: {e}")

    def log(self, msg):
        timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
        full_msg = f"[{timestamp}] {msg}"
        if self.gui_log:
            self.gui_log.append(full_msg)
            QApplication.processEvents()
        if self.log_file:
            try:
                dirname = os.path.dirname(self.log_file)
                if dirname and not os.path.exists(dirname):
                    os.makedirs(dirname, exist_ok=True)
                with open(self.log_file, 'a', encoding='utf-8') as f:
                    f.write(full_msg + '\n')
            except Exception as e:
                if self.gui_log:
                    self.gui_log.append(f"[Ошибка записи лога: {e}]")

logger = Logger()

# ----------------------------------------------------------------------
# Чтение CSV-файлов (только Rigol old, Rigol DHO814, AKIP)
# ----------------------------------------------------------------------
def read_oscilloscope_file(filename, format_type):
    if format_type == 'Rigol old':
        return read_rigol_old(filename)
    elif format_type == 'Rigol DHO814':
        return read_rigol_dho814(filename)
    elif format_type == 'AKIP':
        return read_akip(filename)
    else:
        raise ValueError(f"Неизвестный формат: {format_type}")

def read_rigol_old(filename):
    with open(filename, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    params_line = None
    for line in lines:
        if line.strip():
            parts = line.strip().split(',')
            try:
                float(parts[0])
                float(parts[1])
                params_line = line
                break
            except:
                continue
    if params_line is None:
        raise ValueError("Не найдена строка параметров в формате Rigol old")
    params = params_line.strip().split(',')
    start = float(params[1])
    inc = float(params[2])
    data = []
    for line in lines[lines.index(params_line)+1:]:
        if not line.strip():
            continue
        vals = line.strip().split(',')
        if len(vals) < 2:
            continue
        try:
            data.append([float(v) for v in vals])
        except:
            continue
    if not data:
        raise ValueError("Не найдено числовых данных в файле")
    data = np.array(data)
    t = start + np.arange(data.shape[0]) * inc
    channels = data[:, 1:] if data.shape[1] > 1 else np.zeros((data.shape[0], 1))
    return t, channels

def read_rigol_dho814(filename):
    with open(filename, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    params_line = None
    for i, line in enumerate(lines):
        if 't0' in line.lower() and 'tinc' in line.lower():
            params_line = line
            break
    if params_line is None:
        raise ValueError("Не удалось найти строку с t0 и tInc в заголовке")
    # Добавлен \+ для захвата знака плюс
    t0_match = re.search(r't0\s*=\s*([\d\.\-eE\+]+)', params_line)
    tInc_match = re.search(r'tInc\s*=\s*([\d\.\-eE\+]+)', params_line)
    if not t0_match or not tInc_match:
        raise ValueError("Не удалось извлечь t0 или tInc из строки параметров")
    t0 = float(t0_match.group(1))
    tInc = float(tInc_match.group(1))
    data_lines = []
    start_reading = False
    for line in lines:
        if line is params_line:
            start_reading = True
            continue
        if not start_reading:
            continue
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in line.split(',') if p.strip() != '']
        if not parts:
            continue
        row = []
        for p in parts:
            try:
                val = float(p)
                row.append(val)
            except ValueError:
                # Если число заканчивается на 'e' или 'E', пробуем добавить '+00'
                if p.endswith('e') or p.endswith('E'):
                    p_corrected = p + '+00'
                    try:
                        val = float(p_corrected)
                        row.append(val)
                        continue
                    except:
                        pass
                # Иначе игнорируем это поле
                pass
        if row:
            data_lines.append(row)
    if not data_lines:
        raise ValueError("Не найдено числовых данных в файле")
    # Выравниваем длины строк (если разное количество полей)
    max_len = max(len(row) for row in data_lines)
    data = np.array([row + [0.0]*(max_len - len(row)) for row in data_lines])
    N = data.shape[0]
    t = t0 + np.arange(N) * tInc
    return t, data

def read_akip(filename):
    data = np.loadtxt(filename, delimiter=',', skiprows=1)
    if data.ndim == 1:
        data = data.reshape(-1, 2)
    t = data[:, 0]
    channels = data[:, 1].reshape(-1, 1)
    return t, channels

def interpolate_peak(f_axis, spectrum, idx):
    if idx < 1 or idx >= len(spectrum)-1:
        return f_axis[idx], spectrum[idx]
    x1, y1 = f_axis[idx-1], spectrum[idx-1]
    x2, y2 = f_axis[idx],   spectrum[idx]
    x3, y3 = f_axis[idx+1], spectrum[idx+1]
    denom = (x1 - x2)*(x1 - x3)*(x2 - x3)
    if abs(denom) < 1e-12:
        return x2, y2
    a = (x3*(y2 - y1) + x2*(y1 - y3) + x1*(y3 - y2)) / denom
    b = (x3**2*(y1 - y2) + x2**2*(y3 - y1) + x1**2*(y2 - y3)) / denom
    if abs(a) < 1e-12:
        return x2, y2
    freq = -b / (2*a)
    ampl = a * freq**2 + b * freq + (y1 - a*x1**2 - b*x1)
    if ampl < 0:
        ampl = y2
    if freq < x1 or freq > x3:
        freq, ampl = x2, y2
    return freq, ampl

# ----------------------------------------------------------------------
# Анализ гармоник и ГОСТ
# ----------------------------------------------------------------------
class GOSTAnalyzer:
    def __init__(self, voltage_level=1, U_nom=220):
        self.voltage_level = voltage_level
        self.U_nom = U_nom
        self.gost_table1 = {
            5: [6.0, 4.0, 3.0, 1.5],
            7: [5.0, 3.0, 2.5, 1.0],
            11: [3.5, 2.0, 2.0, 1.0],
            13: [3.0, 2.0, 1.5, 0.7],
            17: [2.0, 1.5, 1.0, 0.5],
            19: [1.5, 1.0, 1.0, 0.4],
            23: [1.5, 1.0, 1.0, 0.4],
            25: [1.5, 1.0, 1.0, 0.4],
            '>25': [1.5, 1.0, 1.0, 0.4]
        }
        self.gost_table2 = {
            3: [5.0, 3.0, 3.0, 1.5],
            9: [1.5, 1.0, 1.0, 0.4],
            15: [0.3, 0.3, 0.3, 0.2],
            21: [0.2, 0.2, 0.2, 0.2],
            '>21': [0.2, 0.2, 0.2, 0.2]
        }
        self.gost_table3 = {
            2: [2.0, 1.5, 1.0, 0.5],
            4: [1.0, 0.7, 0.5, 0.3],
            6: [0.5, 0.3, 0.3, 0.2],
            8: [0.5, 0.3, 0.3, 0.2],
            10: [0.5, 0.3, 0.3, 0.2],
            12: [0.2, 0.2, 0.2, 0.2],
            '>12': [0.2, 0.2, 0.2, 0.2]
        }
        self.gost_table4 = [8.0, 5.0, 4.0, 2.0]
        self.gost_table5 = [12.0, 8.0, 6.0, 3.0]

    def get_gost_limit(self, harmonic_num):
        if harmonic_num == 1:
            return None, 'Основная'
        if harmonic_num % 2 == 0:
            if harmonic_num <= 12:
                key = harmonic_num
                if key in self.gost_table3:
                    return self.gost_table3[key][self.voltage_level-1], 'Таблица 3 (четные)'
            else:
                return self.gost_table3['>12'][self.voltage_level-1], 'Таблица 3 (четные >12)'
        else:
            if harmonic_num % 3 == 0:
                if harmonic_num <= 21:
                    key = harmonic_num
                    if key in self.gost_table2:
                        return self.gost_table2[key][self.voltage_level-1], 'Таблица 2 (нечетные кратные 3)'
                else:
                    return self.gost_table2['>21'][self.voltage_level-1], 'Таблица 2 (нечетные кратные 3 >21)'
            else:
                if harmonic_num <= 25:
                    key = harmonic_num
                    if key in self.gost_table1:
                        return self.gost_table1[key][self.voltage_level-1], 'Таблица 1 (нечетные не кратные 3)'
                else:
                    return self.gost_table1['>25'][self.voltage_level-1], 'Таблица 1 (нечетные не кратные 3 >25)'
        return None, 'Не нормируется'

    def analyze_voltage(self, signal, fs, f0=None):
        if not np.all(np.isfinite(signal)):
            raise ValueError("Сигнал содержит нечисловые значения (nan или inf)")
        N = len(signal)
        signal = signal - np.mean(signal)
        window = windows.hamming(N)
        win_signal = signal * window
        U_fft = fft(win_signal)
        U_fft_norm = U_fft / np.sum(window)
        f_axis = fs * np.arange(N//2 + 1) / N
        U_amp = np.abs(U_fft_norm[:N//2+1])
        U_amp[1:-1] *= 2

        if f0 is None:
            search = np.where((f_axis >= 45) & (f_axis <= 55))[0]
            if len(search) == 0:
                raise ValueError("Не найдено пика в диапазоне 45-55 Гц")
            idx0 = search[np.argmax(U_amp[search])]
            f0, fund_amp = interpolate_peak(f_axis, U_amp, idx0)
        else:
            idx0 = np.argmin(np.abs(f_axis - f0))
            f0, fund_amp = interpolate_peak(f_axis, U_amp, idx0)

        max_harm = 40
        harm_freqs = np.zeros(max_harm)
        harm_amps = np.zeros(max_harm)
        harm_phases = np.zeros(max_harm)
        for k in range(1, max_harm+1):
            target = k * f0
            idx = np.argmin(np.abs(f_axis - target))
            search_range = max(0, idx-3), min(len(f_axis)-1, idx+3)
            local_idx = search_range[0] + np.argmax(U_amp[search_range[0]:search_range[1]+1])
            freq, amp = interpolate_peak(f_axis, U_amp, local_idx)
            harm_freqs[k-1] = freq
            harm_amps[k-1] = amp
            harm_phases[k-1] = np.angle(U_fft_norm[local_idx], deg=True)

        rel_amps = 100 * harm_amps / harm_amps[0]
        U_rms = np.sqrt(np.mean(signal**2))
        U_dev = (U_rms - self.U_nom) / self.U_nom * 100
        U_status = 'СООТВЕТСТВУЕТ' if abs(U_dev) <= 10 else 'НЕ СООТВЕТСТВУЕТ'

        gost_limits = np.full(max_harm, np.nan)
        gost_tables = [''] * max_harm
        status_harm = [''] * max_harm
        violations = []
        warnings = []
        for k in range(1, max_harm+1):
            limit, table = self.get_gost_limit(k)
            gost_limits[k-1] = limit
            gost_tables[k-1] = table if limit is not None else 'Не нормируется'
            if k == 1:
                status_harm[k-1] = 'Основная'
            elif limit is None:
                status_harm[k-1] = 'Не нормируется'
            else:
                if rel_amps[k-1] > limit:
                    status_harm[k-1] = 'НАРУШЕНИЕ'
                    violations.append((k, rel_amps[k-1], limit, rel_amps[k-1]-limit))
                elif rel_amps[k-1] > 0.8 * limit:
                    status_harm[k-1] = 'ВНИМАНИЕ'
                    warnings.append((k, rel_amps[k-1], limit, limit-rel_amps[k-1]))
                else:
                    status_harm[k-1] = 'СООТВЕТСТВУЕТ'

        harm_power = np.sum(harm_amps[1:]**2)
        Ku = 100 * np.sqrt(harm_power) / harm_amps[0]
        Ku_limit_95 = self.gost_table4[self.voltage_level-1]
        Ku_limit_100 = self.gost_table5[self.voltage_level-1]
        if Ku > Ku_limit_100:
            Ku_status = 'НЕ СООТВЕТСТВУЕТ'
        elif Ku > Ku_limit_95:
            Ku_status = 'ПРЕДУПРЕЖДЕНИЕ'
        else:
            Ku_status = 'СООТВЕТСТВУЕТ'

        freq_dev = f0 - 50
        if abs(freq_dev) <= 0.2:
            freq_status = 'НОРМАЛЬНО ДОПУСТИМОЕ'
        elif abs(freq_dev) <= 0.4:
            freq_status = 'ПРЕДЕЛЬНО ДОПУСТИМОЕ'
        else:
            freq_status = 'НЕДОПУСТИМОЕ'

        U_max = np.max(signal)
        U_min = np.min(signal)
        U_peak = np.max(np.abs(signal))
        crest_factor_U = U_peak / U_rms if U_rms > 0 else 0

        result = {
            'f0': f0,
            'freq_dev': freq_dev,
            'freq_status': freq_status,
            'U_rms': U_rms,
            'U_dev': U_dev,
            'U_status': U_status,
            'U_max': U_max,
            'U_min': U_min,
            'U_peak': U_peak,
            'crest_factor_U': crest_factor_U,
            'harm_freqs': harm_freqs,
            'harm_amps': harm_amps,
            'harm_phases': harm_phases,
            'rel_amps': rel_amps,
            'gost_limits': gost_limits,
            'gost_tables': gost_tables,
            'status_harm': status_harm,
            'violations': violations,
            'warnings': warnings,
            'Ku': Ku,
            'Ku_limit_95': Ku_limit_95,
            'Ku_limit_100': Ku_limit_100,
            'Ku_status': Ku_status,
            'signal': signal,
            'f_axis': f_axis,
            'U_amp': U_amp,
            'Fs': fs,
            'U_nom': self.U_nom,
        }
        return result

    def analyze_current(self, signal, fs, f0):
        if not np.all(np.isfinite(signal)):
            raise ValueError("Сигнал тока содержит нечисловые значения")
        N = len(signal)
        signal = signal - np.mean(signal)
        window = windows.hamming(N)
        win_signal = signal * window
        I_fft = fft(win_signal)
        I_fft_norm = I_fft / np.sum(window)
        f_axis = fs * np.arange(N//2 + 1) / N
        I_amp = np.abs(I_fft_norm[:N//2+1])
        I_amp[1:-1] *= 2

        max_harm = 40
        harm_amps = np.zeros(max_harm)
        harm_phases = np.zeros(max_harm)
        for k in range(1, max_harm+1):
            target = k * f0
            idx = np.argmin(np.abs(f_axis - target))
            search_range = max(0, idx-3), min(len(f_axis)-1, idx+3)
            local_idx = search_range[0] + np.argmax(I_amp[search_range[0]:search_range[1]+1])
            freq, amp = interpolate_peak(f_axis, I_amp, local_idx)
            harm_amps[k-1] = amp
            harm_phases[k-1] = np.angle(I_fft_norm[local_idx], deg=True)

        I_rms = np.sqrt(np.mean(signal**2))
        I_peak = np.max(np.abs(signal))
        crest_factor_I = I_peak / I_rms if I_rms > 0 else 0
        harm_power = np.sum(harm_amps[1:]**2)
        THD_I = 100 * np.sqrt(harm_power) / harm_amps[0] if harm_amps[0] > 0 else 0

        result = {
            'harm_amps': harm_amps,
            'harm_phases': harm_phases,
            'I_rms': I_rms,
            'I_peak': I_peak,
            'crest_factor_I': crest_factor_I,
            'THD_I': THD_I,
            'signal': signal,
            'f_axis': f_axis,
            'I_amp': I_amp,
            'Fs': fs,
        }
        return result

# ----------------------------------------------------------------------
# Мощность и длительный анализ
# ----------------------------------------------------------------------
class PowerAnalyzer:
    @staticmethod
    def compute_power(U_harm_amps, I_harm_amps, U_harm_phases, I_harm_phases, U_rms, I_rms):
        delta_phase = U_harm_phases - I_harm_phases
        delta_rad = np.deg2rad(delta_phase)
        cos_phi = np.cos(delta_rad)
        P = np.sum((U_harm_amps * I_harm_amps) / 2 * cos_phi)
        S = U_rms * I_rms
        if S > 0:
            cos_phi_total = P / S
            Q = np.sqrt(S**2 - P**2) if S >= P else 0
        else:
            cos_phi_total = 0
            Q = 0
        if P < 0:
            P = -P
            Q = -Q
        return {
            'P': P,
            'Q': Q,
            'S': S,
            'cos_phi_total': cos_phi_total,
            'cos_phi_harm': cos_phi,
            'delta_phase_deg': delta_phase,
        }

class LongTermAnalyzer:
    @staticmethod
    def analyze(signal, fs, U_nom=230, freq_estimate=50, rms_tolerance=10.0):
        N = len(signal)
        if N < 1000:
            f0 = freq_estimate
        else:
            window = windows.hamming(N)
            win_sig = signal * window
            S = np.abs(np.fft.rfft(win_sig))
            freqs = np.fft.rfftfreq(N, 1/fs)
            idx = np.where((freqs >= 45) & (freqs <= 55))[0]
            if len(idx) == 0:
                f0 = freq_estimate
            else:
                f0 = freqs[idx[np.argmax(S[idx])]]
        T = 1.0 / f0
        samples_per_period = int(fs * T)
        if samples_per_period < 10:
            samples_per_period = int(fs * 0.02)
        num_periods = N // samples_per_period
        if num_periods == 0:
            return {'error': 'Слишком мало данных для анализа по периодам'}

        rms_values = []
        peak_values = []
        violations = 0
        U_min_rms = float('inf')
        U_max_rms = -float('inf')
        peak_abs_max = 0.0

        for i in range(num_periods):
            start = i * samples_per_period
            end = start + samples_per_period
            segment = signal[start:end]
            seg_mean = np.mean(segment)
            seg_ac = segment - seg_mean
            rms = np.sqrt(np.mean(seg_ac**2))
            rms_values.append(rms)
            peak_abs = np.max(np.abs(seg_ac))
            peak_values.append(peak_abs)
            if rms < U_min_rms:
                U_min_rms = rms
            if rms > U_max_rms:
                U_max_rms = rms
            if peak_abs > peak_abs_max:
                peak_abs_max = peak_abs
            if abs(rms - U_nom) / U_nom * 100 > rms_tolerance:
                violations += 1

        delta = peak_abs_max - U_min_rms
        rms_mean = np.mean(rms_values)
        rms_std = np.std(rms_values)

        result = {
            'f0': f0,
            'num_periods': num_periods,
            'U_min_rms': U_min_rms,
            'U_max_rms': U_max_rms,
            'U_mean_rms': rms_mean,
            'U_std_rms': rms_std,
            'peak_abs_max': peak_abs_max,
            'delta': delta,
            'violations': violations,
            'rms_values': rms_values,
            'peak_values': peak_values,
        }
        return result

# ----------------------------------------------------------------------
# Генерация отчётов (9 графиков, 300 dpi)
# ----------------------------------------------------------------------
class ReportGenerator:
    def __init__(self, output_folder, base_filename):
        self.output_folder = output_folder
        self.base_filename = base_filename
        self.excel_path = os.path.join(output_folder, f"Full_results_{base_filename}.xlsx")
        self.word_path = os.path.join(output_folder, f"Full_report_{base_filename}.docx")
        self.image_paths = []
        self.long_term_results = None
        self.summary_data = None
        self.three_phase_mode = False

    def add_image(self, fig, name, dpi=300):
        path = os.path.join(self.output_folder, f"{name}_{self.base_filename}.png")
        fig.savefig(path, dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        self.image_paths.append(path)
        return path
    
    def set_summary_data(self, summary_data):
        self.summary_data = summary_data
        if summary_data is not None and summary_data.get('num_phases', 0) == 3:
            self.three_phase_mode = True

    # ---------- Вспомогательные функции для подписей ----------
    def _add_bar_labels(self, ax, bars, fmt="{:.2f}", fontsize=8, offset=0.02):
        """Добавляет подписи над столбцами с автоматическим смещением."""
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                ax.text(bar.get_x() + bar.get_width()/2., height + offset,
                        fmt.format(height), ha='center', va='bottom', fontsize=fontsize)

    def _add_stem_labels(self, ax, x, y, fmt="{:.2f}", fontsize=8, offset=0.02):
        """Добавляет подписи над точками stem-графика."""
        for xi, yi in zip(x, y):
            if yi > 0:
                ax.text(xi, yi + offset, fmt.format(yi), ha='center', va='bottom', fontsize=fontsize)

    # ---------- 5 отдельных графиков для одной фазы (только напряжение) ----------
    def generate_single_phase_plots_separate(self, v_res, phase_label="Фаза A", show_labels=True):
        """Создаёт 5 отдельных изображений: сигнал, спектр, гармоники 1-40, гармоники 2-40, RMS по периодам."""
        f0 = v_res['f0']
        fs = v_res['Fs']
        t = np.arange(len(v_res['signal'])) / fs
        paths = []

        # 1. Сигнал
        fig1, ax1 = plt.subplots(figsize=(14.8, 8))
        ax1.plot(t, v_res['signal'], 'b-', linewidth=0.8)
        ax1.set_xlabel('Время, с')
        ax1.set_ylabel('Напряжение, В')
        ax1.set_title(f'Сигнал {phase_label}')
        ax1.grid(True)
        paths.append(self.add_image(fig1, f"signal_{phase_label}"))

        # 2. Спектр
        fig2, ax2 = plt.subplots(figsize=(14.8, 8))
        max_freq_show = min(fs/2, 45*f0)
        idx = v_res['f_axis'] <= max_freq_show
        ax2.stem(v_res['f_axis'][idx], v_res['U_amp'][idx], basefmt=" ")
        self._add_stem_labels(ax2, v_res['f_axis'][idx], v_res['U_amp'][idx], fmt="{:.2f}", fontsize=7, offset=0.5)
        ax2.set_xlabel('Частота, Гц')
        ax2.set_ylabel('Амплитуда, В')
        ax2.set_title(f'Амплитудный спектр {phase_label}')
        ax2.grid(True)
        paths.append(self.add_image(fig2, f"spectrum_{phase_label}"))

        # 3. Гармоники 1..40
        fig3, ax3 = plt.subplots(figsize=(14.8, 8))
        harm = np.arange(1, 41)
        rel = v_res['rel_amps']
        limits = v_res['gost_limits']
        colors = []
        for k in range(1,41):
            if k == 1:
                colors.append('green')
            elif np.isnan(limits[k-1]):
                colors.append('gray')
            elif rel[k-1] > limits[k-1]:
                colors.append('red')
            elif rel[k-1] > 0.8*limits[k-1]:
                colors.append('orange')
            else:
                colors.append('blue')
        bars = ax3.bar(harm, rel, color=colors)
        self._add_bar_labels(ax3, bars, fmt="{:.2f}", fontsize=7, offset=0.5)
        for k in range(2,41):
            if not np.isnan(limits[k-1]):
                ax3.plot([k-0.4, k+0.4], [limits[k-1], limits[k-1]], 'r--', linewidth=2)
        ax3.set_xlabel('Номер гармоники')
        ax3.set_ylabel('Относительная амплитуда, %')
        ax3.set_title(f'Гармонический состав (1..40), Ku={v_res["Ku"]:.2f}%  {phase_label}')
        ax3.grid(True)
        ax3.set_xlim(0, 41)
        paths.append(self.add_image(fig3, f"harmonics_1_40_{phase_label}"))

        # 4. Гармоники 2..40 (увеличенный)
        fig4, ax4 = plt.subplots(figsize=(14.8, 8))
        harm2 = np.arange(2,41)
        rel2 = rel[1:]
        lim2 = limits[1:]
        colors2 = []
        for i, k in enumerate(harm2):
            if np.isnan(lim2[i]):
                colors2.append('gray')
            elif rel2[i] > lim2[i]:
                colors2.append('red')
            elif rel2[i] > 0.8*lim2[i]:
                colors2.append('orange')
            else:
                colors2.append('blue')
        bars2 = ax4.bar(harm2, rel2, color=colors2)
        self._add_bar_labels(ax4, bars2, fmt="{:.2f}", fontsize=7, offset=0.3)
        for i, k in enumerate(harm2):
            if not np.isnan(lim2[i]):
                ax4.plot([k-0.4, k+0.4], [lim2[i], lim2[i]], 'r--', linewidth=2)
        ax4.set_xlabel('Номер гармоники (2..40)')
        ax4.set_ylabel('Относительная амплитуда, %')
        ax4.set_title(f'Гармоники 2..40 с пределами по ГОСТ  {phase_label}')
        ax4.grid(True)
        ax4.set_xlim(1.5, 40.5)
        finite_rel = rel2[np.isfinite(rel2)]
        if len(finite_rel) > 0:
            ax4.set_ylim(0, max(finite_rel)*1.2 + 0.5)
        else:
            ax4.set_ylim(0, 1)
        paths.append(self.add_image(fig4, f"harmonics_2_40_{phase_label}"))

        # 5. RMS по периодам с пределами ±3%
        fig5, ax5 = plt.subplots(figsize=(14.8, 8))
        period_samples = int(round(fs / f0))
        num_periods = len(v_res['signal']) // period_samples
        if num_periods >= 2:
            Urms_period = []
            for p in range(num_periods):
                seg = v_res['signal'][p*period_samples:(p+1)*period_samples]
                Urms_period.append(np.sqrt(np.mean(seg**2)))
            ax5.plot(range(1, num_periods+1), Urms_period, 'b-o', linewidth=1, markersize=6)
            if show_labels:
                for i, val in enumerate(Urms_period):
                    ax5.text(i+1, val + 0.5, f"{val:.1f}", ha='center', va='bottom', fontsize=8)
            ax5.axhline(v_res['U_nom'], color='r', linestyle='--', label='Номинал')
            ax5.axhline(v_res['U_nom']*0.97, color='g', linestyle='--', label='-3%')
            ax5.axhline(v_res['U_nom']*1.03, color='g', linestyle='--', label='+3%')
            ax5.set_xlabel('Номер периода')
            ax5.set_ylabel('Действующее напряжение, В')
            ax5.set_title(f'RMS по периодам (±3%)  {phase_label}')
            ax5.grid(True)
            ax5.legend()
        else:
            ax5.text(0.5, 0.5, 'Недостаточно периодов', ha='center', va='center', transform=ax5.transAxes)
            ax5.axis('off')
        paths.append(self.add_image(fig5, f"rms_periods_{phase_label}"))

        return paths

    # ---------- Общий составной график (5 подграфиков) для одной фазы (только напряжение) ----------
    def generate_single_phase_plots_combined(self, v_res, phase_label="Фаза A", show_labels=True):
        """Составной график с 5 подграфиками (сигнал, спектр, гармоники 1-40, 2-40, RMS)."""
        f0 = v_res['f0']
        fs = v_res['Fs']
        t = np.arange(len(v_res['signal'])) / fs

        fig, axes = plt.subplots(2, 3, figsize=(14.8, 8))
        ax1, ax2, ax3, ax4, ax5 = axes[0,0], axes[0,1], axes[0,2], axes[1,0], axes[1,1]
        axes[1,2].axis('off')

        # 1. Сигнал
        ax1.plot(t, v_res['signal'], 'b-', linewidth=0.8)
        ax1.set_xlabel('Время, с')
        ax1.set_ylabel('Напряжение, В')
        ax1.set_title(f'Сигнал {phase_label}')
        ax1.grid(True)

        # 2. Спектр
        max_freq_show = min(fs/2, 45*f0)
        idx = v_res['f_axis'] <= max_freq_show
        ax2.stem(v_res['f_axis'][idx], v_res['U_amp'][idx], basefmt=" ")
        ax2.set_xlabel('Частота, Гц')
        ax2.set_ylabel('Амплитуда, В')
        ax2.set_title('Амплитудный спектр')
        ax2.grid(True)

        # 3. Гармоники 1..40
        harm = np.arange(1, 41)
        rel = v_res['rel_amps']
        limits = v_res['gost_limits']
        colors = []
        for k in range(1,41):
            if k == 1:
                colors.append('green')
            elif np.isnan(limits[k-1]):
                colors.append('gray')
            elif rel[k-1] > limits[k-1]:
                colors.append('red')
            elif rel[k-1] > 0.8*limits[k-1]:
                colors.append('orange')
            else:
                colors.append('blue')
        ax3.bar(harm, rel, color=colors)
        for k in range(2,41):
            if not np.isnan(limits[k-1]):
                ax3.plot([k-0.4, k+0.4], [limits[k-1], limits[k-1]], 'r--', linewidth=2)
        ax3.set_xlabel('Номер гармоники')
        ax3.set_ylabel('Относительная амплитуда, %')
        ax3.set_title(f'Гармоники 1..40, Ku={v_res["Ku"]:.2f}%')
        ax3.grid(True)
        ax3.set_xlim(0, 41)

        # 4. Гармоники 2..40
        harm2 = np.arange(2,41)
        rel2 = rel[1:]
        lim2 = limits[1:]
        colors2 = []
        for i, k in enumerate(harm2):
            if np.isnan(lim2[i]):
                colors2.append('gray')
            elif rel2[i] > lim2[i]:
                colors2.append('red')
            elif rel2[i] > 0.8*lim2[i]:
                colors2.append('orange')
            else:
                colors2.append('blue')
        ax4.bar(harm2, rel2, color=colors2)
        for i, k in enumerate(harm2):
            if not np.isnan(lim2[i]):
                ax4.plot([k-0.4, k+0.4], [lim2[i], lim2[i]], 'r--', linewidth=2)
        ax4.set_xlabel('Номер гармоники (2..40)')
        ax4.set_ylabel('Относительная амплитуда, %')
        ax4.set_title('Гармоники 2..40 с пределами ГОСТ')
        ax4.grid(True)
        ax4.set_xlim(1.5, 40.5)
        finite_rel = rel2[np.isfinite(rel2)]
        if len(finite_rel) > 0:
            ax4.set_ylim(0, max(finite_rel)*1.2 + 0.5)
        else:
            ax4.set_ylim(0, 1)

        # 5. RMS по периодам
        period_samples = int(round(fs / f0))
        num_periods = len(v_res['signal']) // period_samples
        if num_periods >= 2:
            Urms_period = []
            for p in range(num_periods):
                seg = v_res['signal'][p*period_samples:(p+1)*period_samples]
                Urms_period.append(np.sqrt(np.mean(seg**2)))
            ax5.plot(range(1, num_periods+1), Urms_period, 'b-o', linewidth=1, markersize=6)
            if show_labels:
                for i, val in enumerate(Urms_period):
                    ax5.text(i+1, val + 0.5, f"{val:.1f}", ha='center', va='bottom', fontsize=8)
            ax5.axhline(v_res['U_nom'], color='r', linestyle='--', label='Номинал')
            ax5.axhline(v_res['U_nom']*0.97, color='g', linestyle='--', label='-3%')
            ax5.axhline(v_res['U_nom']*1.03, color='g', linestyle='--', label='+3%')
            ax5.set_xlabel('Номер периода')
            ax5.set_ylabel('Действующее напряжение, В')
            ax5.set_title('RMS по периодам (±3%)')
            ax5.grid(True)
            ax5.legend()
        else:
            ax5.text(0.5, 0.5, 'Недостаточно периодов', ha='center', va='center', transform=ax5.transAxes)
            ax5.axis('off')

        fig.suptitle(f'Анализ гармоник, {phase_label}, f0={f0:.3f} Гц')
        plt.tight_layout()
        return self.add_image(fig, f"combined_analysis_{phase_label}")

    # ---------- Составной график (6 подграфиков) для пары U+I ----------
    def generate_ui_phase_plots(self, v_res, i_res, power_res, phase_label="Фаза A"):
        """6 подграфиков: сигналы U/I, спектры U/I, круговая диаграмма, cos φ, относительные амплитуды."""
        f0 = v_res['f0']
        fs = v_res['Fs']
        t = np.arange(len(v_res['signal'])) / fs

        U_norm = v_res['signal'] / np.max(np.abs(v_res['signal'])) if np.max(np.abs(v_res['signal'])) > 0 else v_res['signal']
        I_norm = i_res['signal'] / np.max(np.abs(i_res['signal'])) if np.max(np.abs(i_res['signal'])) > 0 else i_res['signal']

        fig, axes = plt.subplots(2, 3, figsize=(14.8, 8))
        ax1, ax2, ax3, ax4, ax5, ax6 = axes.flatten()

        # 1. Временные сигналы
        ax1.plot(t, U_norm, 'r-', label='U (норм.)', linewidth=0.8)
        ax1.plot(t, I_norm, 'b-', label='I (норм.)', linewidth=0.8)
        ax1.set_xlabel('Время, с')
        ax1.set_ylabel('Нормированные значения')
        ax1.set_title('Напряжение и ток (первые 3 периода)')
        ax1.legend()
        ax1.grid(True)

        # 2. Спектр U (2..40)
        harm = np.arange(2,41)
        ax2.bar(harm, v_res['harm_amps'][1:], color='red', alpha=0.7)
        ax2.set_xlabel('Номер гармоники (2..40)')
        ax2.set_ylabel('Амплитуда, В')
        ax2.set_title(f'Спектр U, U1={v_res["harm_amps"][0]:.2f} В')
        ax2.grid(True)

        # 3. Спектр I (2..40)
        ax3.bar(harm, i_res['harm_amps'][1:], color='blue', alpha=0.7)
        ax3.set_xlabel('Номер гармоники (2..40)')
        ax3.set_ylabel('Амплитуда, А')
        ax3.set_title(f'Спектр I, I1={i_res["harm_amps"][0]:.2f} А')
        ax3.grid(True)

        # 4. Круговая диаграмма мощностей
        if power_res is not None:
            Q_abs = abs(power_res['Q'])
            if Q_abs < 1e-6:
                Q_abs = 0
            pie_data = [power_res['P'], Q_abs]
            if sum(pie_data) == 0:
                pie_data = [1, 0]
            ax4.pie(pie_data, labels=['Активная P', 'Реактивная Q'], autopct='%1.1f%%')
            ax4.set_title(f"P = {power_res['P']:.1f} Вт, Q = {power_res['Q']:.1f} ВАр\ncos φ = {power_res['cos_phi_total']:.4f}")
        else:
            ax4.text(0.5, 0.5, 'Нет данных мощности', ha='center', va='center', transform=ax4.transAxes)
            ax4.axis('off')

        # 5. Коэффициент мощности по гармоникам
        if power_res is not None:
            cos_phi = power_res['cos_phi_harm']
            ax5.bar(range(1,41), cos_phi, color='0.6')
            ax5.set_xlabel('Номер гармоники')
            ax5.set_ylabel('cos φ')
            ax5.set_title('Коэффициент мощности по гармоникам')
            ax5.set_ylim(-1.1, 1.1)
            ax5.grid(True)
        else:
            ax5.text(0.5, 0.5, 'Нет данных', ha='center', va='center', transform=ax5.transAxes)
            ax5.axis('off')

        # 6. Относительные амплитуды U и I
        rel_U = v_res['rel_amps'][1:]
        rel_I = 100 * i_res['harm_amps'][1:] / i_res['harm_amps'][0]
        width = 0.35
        ax6.bar(harm - width/2, rel_U, width, label='U', alpha=0.7, color='red')
        ax6.bar(harm + width/2, rel_I, width, label='I', alpha=0.7, color='blue')
        ax6.set_xlabel('Номер гармоники (2..40)')
        ax6.set_ylabel('Относительная амплитуда, %')
        ax6.set_title('Относительные амплитуды U и I (без 1-й)')
        ax6.legend()
        ax6.grid(True)

        fig.suptitle(f'Анализ напряжения и тока, {phase_label}, f0={f0:.3f} Гц')
        plt.tight_layout()
        return self.add_image(fig, f"UI_analysis_{phase_label}")

    # ---------- Отдельные графики для каждого подграфика (U+I) ----------
    def generate_ui_phase_plots_separate(self, v_res, i_res, power_res, phase_label="Фаза A"):
        """Создаёт 6 отдельных изображений для каждого подграфика из UI_analysis."""
        f0 = v_res['f0']
        fs = v_res['Fs']
        t = np.arange(len(v_res['signal'])) / fs
        paths = []

        # 1. Временные сигналы (нормированные)
        fig1, ax1 = plt.subplots(figsize=(14.8, 8))
        U_norm = v_res['signal'] / np.max(np.abs(v_res['signal'])) if np.max(np.abs(v_res['signal'])) > 0 else v_res['signal']
        I_norm = i_res['signal'] / np.max(np.abs(i_res['signal'])) if np.max(np.abs(i_res['signal'])) > 0 else i_res['signal']
        ax1.plot(t, U_norm, 'r-', label='U (норм.)', linewidth=0.8)
        ax1.plot(t, I_norm, 'b-', label='I (норм.)', linewidth=0.8)
        ax1.set_xlabel('Время, с')
        ax1.set_ylabel('Нормированные значения')
        ax1.set_title(f'Напряжение и ток (нормированные) {phase_label}')
        ax1.legend()
        ax1.grid(True)
        paths.append(self.add_image(fig1, f"signal_{phase_label}"))

        # 2. Спектр U (2..40) с подписями
        fig2, ax2 = plt.subplots(figsize=(14.8, 8))
        harm = np.arange(2,41)
        amps_U = v_res['harm_amps'][1:]
        bars = ax2.bar(harm, amps_U, color='red', alpha=0.7)
        self._add_bar_labels(ax2, bars, fmt="{:.2f}", fontsize=8, offset=0.02*max(amps_U) if max(amps_U)>0 else 0.1)
        ax2.set_xlabel('Номер гармоники (2..40)')
        ax2.set_ylabel('Амплитуда, В')
        ax2.set_title(f'Спектр U, U1={v_res["harm_amps"][0]:.2f} В  {phase_label}')
        ax2.grid(True)
        paths.append(self.add_image(fig2, f"spectrum_U_{phase_label}"))

        # 3. Спектр I (2..40) с подписями
        fig3, ax3 = plt.subplots(figsize=(14.8, 8))
        amps_I = i_res['harm_amps'][1:]
        bars = ax3.bar(harm, amps_I, color='blue', alpha=0.7)
        self._add_bar_labels(ax3, bars, fmt="{:.2f}", fontsize=8, offset=0.02*max(amps_I) if max(amps_I)>0 else 0.1)
        ax3.set_xlabel('Номер гармоники (2..40)')
        ax3.set_ylabel('Амплитуда, А')
        ax3.set_title(f'Спектр I, I1={i_res["harm_amps"][0]:.2f} А  {phase_label}')
        ax3.grid(True)
        paths.append(self.add_image(fig3, f"spectrum_I_{phase_label}"))

        # 4. Круговая диаграмма мощностей
        fig4, ax4 = plt.subplots(figsize=(8, 8))  # квадратная для круговой
        if power_res is not None:
            Q_abs = abs(power_res['Q'])
            if Q_abs < 1e-6:
                Q_abs = 0
            pie_data = [power_res['P'], Q_abs]
            if sum(pie_data) == 0:
                pie_data = [1, 0]
            ax4.pie(pie_data, labels=['Активная P', 'Реактивная Q'], autopct='%1.1f%%')
            ax4.set_title(f"P = {power_res['P']:.2f} Вт, Q = {power_res['Q']:.2f} ВАр\ncos φ = {power_res['cos_phi_total']:.4f}  {phase_label}")
        else:
            ax4.text(0.5, 0.5, 'Нет данных мощности', ha='center', va='center', transform=ax4.transAxes)
            ax4.axis('off')
        paths.append(self.add_image(fig4, f"pie_chart_{phase_label}"))

        # 5. Коэффициент мощности по гармоникам с подписями
        fig5, ax5 = plt.subplots(figsize=(14.8, 8))
        if power_res is not None:
            cos_phi = power_res['cos_phi_harm']
            bars = ax5.bar(range(1,41), cos_phi, color='0.6')
            self._add_bar_labels(ax5, bars, fmt="{:.2f}", fontsize=7, offset=0.03)
            ax5.set_xlabel('Номер гармоники')
            ax5.set_ylabel('cos φ')
            ax5.set_title(f'Коэффициент мощности по гармоникам  {phase_label}')
            ax5.set_ylim(-1.1, 1.1)
            ax5.grid(True)
        else:
            ax5.text(0.5, 0.5, 'Нет данных', ha='center', va='center', transform=ax5.transAxes)
            ax5.axis('off')
        paths.append(self.add_image(fig5, f"cos_phi_{phase_label}"))

        # 6. Относительные амплитуды U и I с подписями
        fig6, ax6 = plt.subplots(figsize=(14.8, 8))
        rel_U = v_res['rel_amps'][1:]
        rel_I = 100 * i_res['harm_amps'][1:] / i_res['harm_amps'][0]
        width = 0.35
        bars_U = ax6.bar(harm - width/2, rel_U, width, label='U', alpha=0.7, color='red')
        bars_I = ax6.bar(harm + width/2, rel_I, width, label='I', alpha=0.7, color='blue')
        
        # Динамический offset и расширение ylim
        max_rel = max(np.max(rel_U), np.max(rel_I)) if len(rel_U) > 0 else 1
        offset_rel = max_rel * 0.03 if max_rel > 0 else 0.5
        self._add_bar_labels(ax6, bars_U, fmt="{:.2f}", fontsize=7, offset=offset_rel)
        self._add_bar_labels(ax6, bars_I, fmt="{:.2f}", fontsize=7, offset=offset_rel)
        ax6.set_ylim(0, max_rel * 1.15)  # запас 15% сверху
        
        ax6.set_xlabel('Номер гармоники (2..40)')
        ax6.set_ylabel('Относительная амплитуда, %')
        ax6.set_title(f'Относительные амплитуды U и I (без 1-й)  {phase_label}')
        ax6.legend()
        ax6.grid(True)
        paths.append(self.add_image(fig6, f"rel_amplitudes_{phase_label}"))

        return paths

    # ---------- Сравнительный график по фазам ----------
    def generate_phase_comparison(self, voltage_results_list, phase_labels):
        fig, ax = plt.subplots(figsize=(14.8, 8))
        colors = ['r', 'g', 'b']
        for i, v_res in enumerate(voltage_results_list):
            rel = v_res['rel_amps']
            ax.plot(range(1,41), rel, color=colors[i % 3], linewidth=1.5, label=phase_labels[i])
        ax.set_xlabel('Номер гармоники')
        ax.set_ylabel('Относительная амплитуда, %')
        ax.set_title('Сравнение относительных амплитуд гармоник по фазам')
        ax.grid(True)
        ax.set_xlim(0, 41)
        ax.legend()
        return self.add_image(fig, "harmonics_compare")

    # ---------- Спектр напряжения (2..40) для всех фаз ----------
    def generate_spectrum_2_40(self, voltage_results_list, phase_labels):
        fig, ax = plt.subplots(figsize=(14.8, 8))
        width = 0.8 / len(voltage_results_list)
        for i, v_res in enumerate(voltage_results_list):
            harm = np.arange(2,41)
            amps = v_res['harm_amps'][1:]
            ax.bar(harm + (i - (len(voltage_results_list)-1)/2)*width, amps, width,
                   label=phase_labels[i], alpha=0.7)
        ax.set_xlabel('Номер гармоники (2..40)')
        ax.set_ylabel('Амплитуда, В')
        ax.set_title('Спектр напряжения (гармоники 2..40)')
        ax.grid(True)
        ax.legend()
        return self.add_image(fig, "spectrum_U_2_40")

    # ---------- Векторная диаграмма для одной фазы ----------
    def generate_vector_diagram(self, v_res, i_res, phase_label):
        fig, ax = plt.subplots(figsize=(8, 8))
        theta_U = np.deg2rad(v_res['harm_phases'][0])
        theta_I = np.deg2rad(i_res['harm_phases'][0])
        length = 10.0
        Ux = length * np.cos(theta_U)
        Uy = length * np.sin(theta_U)
        Ix = length * np.cos(theta_I)
        Iy = length * np.sin(theta_I)
        ax.quiver(0, 0, Ux, Uy, angles='xy', scale_units='xy', scale=1, color='r', width=0.005, label='U')
        ax.quiver(0, 0, Ix, Iy, angles='xy', scale_units='xy', scale=1, color='b', width=0.005, label='I')
        ax.axhline(0, color='k', linestyle='--', linewidth=0.5)
        ax.axvline(0, color='k', linestyle='--', linewidth=0.5)
        ax.axis('equal')
        ax.grid(True)
        ax.set_xlim(-12, 12)
        ax.set_ylim(-12, 12)
        ax.set_xlabel('Действительная ось')
        ax.set_ylabel('Мнимая ось')
        delta = v_res['harm_phases'][0] - i_res['harm_phases'][0]
        ax.set_title(f"Векторная диаграмма (основная гармоника)\n{phase_label}, Δφ = {delta:.1f}°")
        ax.legend(loc='upper right')
        return self.add_image(fig, f"vector_diagram_{phase_label}")

    # ---------- Основной метод генерации всех графиков ----------
    def generate_full_report(self, voltage_results_list, current_results_list=None,
                             power_results_list=None, phase_labels=None):
        if phase_labels is None:
            phase_labels = [f'Фаза {chr(65+i)}' for i in range(len(voltage_results_list))]

        image_paths = []

        # 1. Для каждой фазы
        for idx, v_res in enumerate(voltage_results_list):
            if current_results_list is not None and idx < len(current_results_list):
                # Есть ток – создаём общий составной и отдельные графики
                i_res = current_results_list[idx]
                power_res = power_results_list[idx] if power_results_list else None

                # Общий составной
                img_combined = self.generate_ui_phase_plots(v_res, i_res, power_res, phase_labels[idx])
                image_paths.append(img_combined)

                # Отдельные 6 графиков
                paths_sep = self.generate_ui_phase_plots_separate(v_res, i_res, power_res, phase_labels[idx])
                image_paths.extend(paths_sep)

                # Векторная диаграмма (отдельно)
                img_vec = self.generate_vector_diagram(v_res, i_res, phase_labels[idx])
                image_paths.append(img_vec)

            else:
                # Только напряжение – общий составной + 5 отдельных
                show_labels = not (self.long_term_results is not None and 'error' not in self.long_term_results)
                img_combined = self.generate_single_phase_plots_combined(v_res, phase_labels[idx], show_labels=show_labels)
                image_paths.append(img_combined)

                paths_sep = self.generate_single_phase_plots_separate(v_res, phase_labels[idx], show_labels=show_labels)
                image_paths.extend(paths_sep)

        # 2. Сравнительный график (если более одной фазы)
        if len(voltage_results_list) > 1:
            img_path = self.generate_phase_comparison(voltage_results_list, phase_labels)
            image_paths.append(img_path)

        # 3. Спектр напряжения 2..40 для всех фаз
        if len(voltage_results_list) > 1:
            img_path = self.generate_spectrum_2_40(voltage_results_list, phase_labels)
            image_paths.append(img_path)

        self.image_paths = image_paths
        return image_paths

    def _add_rms_table_to_excel(self, wb):
        if self.long_term_results is None or 'error' in self.long_term_results:
            return
        ws = wb.create_sheet("RMS по периодам")
        ws.append(['Номер периода', 'RMS, В'])
        rms_values = self.long_term_results.get('rms_values', [])
        total = len(rms_values)
        step = self._get_rms_step(total)
        # Выводим значения с шагом
        for i in range(0, total, step):
            ws.append([i+1, rms_values[i]])
        # Если последний период не попал, добавляем его
        if (total - 1) % step != 0:
            ws.append([total, rms_values[-1]])

    def _get_rms_step(self, total_periods):
        """Возвращает шаг для выборки значений RMS, чтобы таблица была не слишком большой."""
        if total_periods <= 200:
            return 1
        # Выводим примерно 20 строк
        step = max(1, total_periods // 20)
        # Округляем до удобных значений
        if step < 10:
            return 5
        elif step < 50:
            return 10
        elif step < 100:
            return 50
        else:
            return 100

    # ---------- Excel (обновлён для многофазного) ----------
    def generate_excel(self, voltage_results_list, current_results_list=None,
                       power_results_list=None, phase_labels=None):
        if phase_labels is None:
            phase_labels = [f'Фаза {chr(65+i)}' for i in range(len(voltage_results_list))]

        wb = openpyxl.Workbook()
        wb.remove(wb.active)

        # Лист "Сводка"
        ws_summary = wb.create_sheet("Сводка")
        headers = ['Фаза', 'f0, Гц', 'Urms, В', 'U откл., %', 'Ku, %', 'Статус Ku',
                   'Нарушений', 'Предупреждений']
        if current_results_list:
            headers.extend(['Irms, А', 'THDI, %', 'cos φ общ.'])
        ws_summary.append(headers)

        for idx, v_res in enumerate(voltage_results_list):
            row = [phase_labels[idx],
                   v_res['f0'],
                   v_res['U_rms'],
                   v_res['U_dev'],
                   v_res['Ku'],
                   v_res['Ku_status'],
                   len(v_res['violations']),
                   len(v_res['warnings'])]
            if current_results_list:
                i_res = current_results_list[idx]
                row.append(i_res['I_rms'])
                row.append(i_res['THD_I'])
                if power_results_list:
                    row.append(power_results_list[idx]['cos_phi_total'])
                else:
                    row.append('')
            ws_summary.append(row)

        # Листы для каждой фазы
        for idx, v_res in enumerate(voltage_results_list):
            phase = phase_labels[idx]
            sheet_name = phase[:31]
            ws = wb.create_sheet(sheet_name)

            # Гармоники напряжения
            ws.append(['№', 'Частота, Гц', 'Амплитуда, В', 'Отн., %',
                       'Предел ГОСТ, %', 'Отклонение, %', 'Статус', 'Таблица ГОСТ'])
            for k in range(1, 41):
                freq = v_res['harm_freqs'][k-1]
                amp = v_res['harm_amps'][k-1]
                rel = v_res['rel_amps'][k-1]
                limit = v_res['gost_limits'][k-1]
                dev = rel - limit if not np.isnan(limit) else ''
                status = v_res['status_harm'][k-1]
                table = v_res['gost_tables'][k-1] if k > 1 else '---'
                ws.append([k, freq, amp, rel, limit if not np.isnan(limit) else '',
                           dev, status, table])

            # Если есть ток – добавляем гармоники тока и cos φ
            if current_results_list:
                i_res = current_results_list[idx]
                ws.append([])
                ws.append(['Гармоники тока'])
                ws.append(['№', 'Частота, Гц', 'Амплитуда, А', 'Отн., %'])
                for k in range(1, 41):
                    freq = v_res['harm_freqs'][k-1]
                    amp = i_res['harm_amps'][k-1]
                    relI = 100 * amp / i_res['harm_amps'][0] if i_res['harm_amps'][0] > 0 else 0
                    ws.append([k, freq, amp, relI])

                if power_results_list:
                    pr = power_results_list[idx]
                    ws.append([])
                    ws.append(['cos φ по гармоникам'])
                    ws.append(['№', 'Частота, Гц', 'Δφ, °', 'cos φ', 'U_ампл, В'])
                    for k in range(1, 41):
                        freq = v_res['harm_freqs'][k-1]
                        delta = pr['delta_phase_deg'][k-1]
                        cosf = pr['cos_phi_harm'][k-1]
                        U_amp = v_res['harm_amps'][k-1]
                        ws.append([k, freq, delta, cosf, U_amp])

        # Лист "Итоги"
        ws_ito = wb.create_sheet("Итоги")
        ws_ito.append(['Параметр', 'Значение'])
        if voltage_results_list:
            v0 = voltage_results_list[0]
            ws_ito.append(['U_nom, В', v0['U_nom']])
            ws_ito.append(['Средняя частота, Гц', np.mean([v['f0'] for v in voltage_results_list])])
            ws_ito.append(['Мин. Urms, В', min([v['U_rms'] for v in voltage_results_list])])
            ws_ito.append(['Макс. Urms, В', max([v['U_rms'] for v in voltage_results_list])])
            ws_ito.append(['Средний Ku, %', np.mean([v['Ku'] for v in voltage_results_list])])
            total_viol = sum([len(v['violations']) for v in voltage_results_list])
            ws_ito.append(['Всего нарушений по гармоникам', total_viol])

        if current_results_list:
            ws_ito.append(['Средний THDI, %', np.mean([i['THD_I'] for i in current_results_list])])
        if power_results_list:
            ws_ito.append(['Средний cos φ общий', np.mean([p['cos_phi_total'] for p in power_results_list])])

        # Длительный анализ (если есть)
        if self.long_term_results is not None and 'error' not in self.long_term_results:
            ws_long = wb.create_sheet("Длительный анализ")
            res = self.long_term_results
            ws_long.append(['Параметр', 'Значение'])
            ws_long.append(['Количество периодов', res['num_periods']])
            ws_long.append(['Минимальное RMS, В', res['U_min_rms']])
            ws_long.append(['Максимальное RMS, В', res['U_max_rms']])
            ws_long.append(['Среднее RMS, В', res['U_mean_rms']])
            ws_long.append(['Станд. отклонение RMS, В', res['U_std_rms']])
            ws_long.append(['Максимальный пик (абс.), В', res['peak_abs_max']])
            ws_long.append(['Дельта (пик - мин.RMS), В', res['delta']])
            ws_long.append(['Количество нарушений (±10%)', res['violations']])

        # --- ЛИСТ "График" с диаграммой (всегда, если есть напряжение) ---
        if voltage_results_list:
            ws_chart = wb.create_sheet("График")
            v_res = voltage_results_list[0]  # берём первую фазу

            # Заголовки: №, Частота, Отн. ампл. %, Предел ГОСТ
            ws_chart.append(['№', 'Частота, Гц', 'Отн. ампл., %', 'Предел ГОСТ, %'])
            for k in range(2, 41):
                freq = v_res['harm_freqs'][k-1]
                rel = v_res['rel_amps'][k-1]
                limit = v_res['gost_limits'][k-1]
                ws_chart.append([k, freq, rel, limit if not np.isnan(limit) else None])

            # Создаём ScatterChart с двумя сериями
            chart = ScatterChart()
            chart.title = "Спектр гармоник (2..40)"
            chart.x_axis.title = 'Частота, Гц'
            chart.y_axis.title = 'Относительная амплитуда, %'
            chart.style = 13

            xvalues = Reference(ws_chart, min_col=2, min_row=2, max_row=40)

            # Серия 1: Отн. ампл.
            yvalues1 = Reference(ws_chart, min_col=3, min_row=2, max_row=40)
            series1 = Series(yvalues1, xvalues, title="Отн. ампл., %")
            series1.marker = openpyxl.chart.marker.Marker('circle')
            series1.graphicalProperties.line.solid = True
            series1.graphicalProperties.line.width = 15000
            chart.series.append(series1)

            # Серия 2: Предел ГОСТ
            yvalues2 = Reference(ws_chart, min_col=4, min_row=2, max_row=40)
            series2 = Series(yvalues2, xvalues, title="Предел ГОСТ")
            series2.marker = openpyxl.chart.marker.Marker('diamond')
            series2.graphicalProperties.line.solid = True
            series2.graphicalProperties.line.width = 15000
            series2.graphicalProperties.line.dashStyle = 'dash'
            chart.series.append(series2)

            ws_chart.add_chart(chart, 'E2')

        # Форматирование всех чисел с двумя знаками после запятой
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, (int, float)):
                        cell.number_format = '0.00'
                    elif isinstance(cell.value, str):
                        try:
                            val = float(cell.value)
                            cell.value = val
                            cell.number_format = '0.00'
                        except:
                            pass

        # Добавить суммарные параметры, если есть
        if self.summary_data is not None:
            sd = self.summary_data
            ws_ito.append([])
            ws_ito.append(['СУММАРНЫЕ ПАРАМЕТРЫ', ''])
            ws_ito.append(['P_total, Вт', sd['P_total']])
            ws_ito.append(['Q_total, ВАр', sd['Q_total']])
            ws_ito.append(['S_total, ВА', sd['S_total']])
            ws_ito.append(['cos φ общий', sd['cos_phi_total']])
            ws_ito.append(['Средний THD_U, %', sd['avg_THD_U']])
            ws_ito.append(['Средний THD_I, %', sd['avg_THD_I']])
            ws_ito.append(['Всего нарушений', sd['total_violations']])
            ws_ito.append(['Всего предупреждений', sd['total_warnings']])
            ws_ito.append(['Общий статус Ku', sd['ku_status_overall']])

                # Если есть длительный анализ – добавляем лист с RMS по периодам
        if self.long_term_results is not None and 'error' not in self.long_term_results:
            self._add_rms_table_to_excel(wb)

        wb.save(self.excel_path)
        return self.excel_path

    # ---------- Word (обновлён для многофазного) ----------
    def generate_word(self, voltage_results_list, current_results_list=None,
                      power_results_list=None, phase_labels=None, image_paths=None):
        from datetime import datetime
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(0.75)
        section.right_margin = Cm(0.75)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.first_line_indent = Cm(1.25)
        style.paragraph_format.line_spacing = 1.5

        # Определяем режимы
        only_voltage = (current_results_list is None or len(current_results_list) == 0) and len(voltage_results_list) > 0
        only_current = (voltage_results_list is None or len(voltage_results_list) == 0) and len(current_results_list) > 0
        both = (voltage_results_list and current_results_list)

        # Заголовок
        if both:
            title = doc.add_heading('ОТЧЁТ ПО АНАЛИЗУ ДВУХКАНАЛЬНОЙ ОСЦИЛЛОГРАММЫ', 0)
        else:
            title = doc.add_heading('ОТЧЁТ ПО АНАЛИЗУ МНОГОФАЗНОЙ ОСЦИЛЛОГРАММЫ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Общая информация ---
        doc.add_heading('Общая информация', level=1)
        if voltage_results_list:
            v0 = voltage_results_list[0]
            doc.add_paragraph(f"Количество фаз: {len(voltage_results_list)}")
            doc.add_paragraph(f"Номинальное напряжение: {v0['U_nom']:.1f} В")
            doc.add_paragraph(f"Частота дискретизации: {v0['Fs']:.2f} Гц")
            doc.add_paragraph(f"Количество отсчётов: {len(v0['signal'])}")
            if both:
                doc.add_paragraph(f"Масштаб напряжения: x1.00")
                doc.add_paragraph(f"Масштаб тока: x1694.00")
                doc.add_paragraph(f"Инверсия тока: нет")
                doc.add_paragraph(f"Канал напряжения: 1, канал тока: 2")
        elif only_current:
            i0 = current_results_list[0]
            doc.add_paragraph(f"Количество фаз: {len(current_results_list)}")
            doc.add_paragraph(f"Частота дискретизации: {i0['Fs']:.2f} Гц")
            doc.add_paragraph(f"Количество отсчётов: {len(i0['signal'])}")
            doc.add_paragraph(f"Масштаб тока: x1694.00")
            doc.add_paragraph(f"Инверсия тока: нет")

        # --- Сводные параметры по фазам (средние значения) ---
        if len(voltage_results_list) > 1:
            doc.add_heading('Сводные параметры по фазам (средние)', level=1)
            avg_freq = np.mean([v['f0'] for v in voltage_results_list])
            avg_rms = np.mean([v['U_rms'] for v in voltage_results_list])
            avg_dev = np.mean([v['U_dev'] for v in voltage_results_list])
            avg_ku = np.mean([v['Ku'] for v in voltage_results_list])
            min_ku = min([v['Ku'] for v in voltage_results_list])
            max_ku = max([v['Ku'] for v in voltage_results_list])
            total_viol = sum([len(v['violations']) for v in voltage_results_list])
            total_warn = sum([len(v['warnings']) for v in voltage_results_list])
            ku_statuses = [v['Ku_status'] for v in voltage_results_list]
            ku_ok = all(s == 'СООТВЕТСТВУЕТ' for s in ku_statuses)

            doc.add_paragraph(f"Средняя частота: {avg_freq:.3f} Гц")
            doc.add_paragraph(f"Среднее действующее напряжение (RMS): {avg_rms:.1f} В")
            doc.add_paragraph(f"Среднее отклонение напряжения: {avg_dev:+.2f}%")
            doc.add_paragraph(f"Средний Ku (THD): {avg_ku:.2f}% (мин: {min_ku:.2f}%, макс: {max_ku:.2f}%)")
            doc.add_paragraph(f"Всего нарушений по гармоникам: {total_viol}")
            doc.add_paragraph(f"Всего предупреждений: {total_warn}")
            if ku_ok:
                doc.add_paragraph("Ku по всем фазам: СООТВЕТСТВУЕТ")
            else:
                doc.add_paragraph("Ku по некоторым фазам: НЕ СООТВЕТСТВУЕТ")

            if both:
                avg_thdi = np.mean([i['THD_I'] for i in current_results_list])
                doc.add_paragraph(f"Средний THDI: {avg_thdi:.2f}%")
                if power_results_list:
                    avg_cos = np.mean([p['cos_phi_total'] for p in power_results_list])
                    doc.add_paragraph(f"Средний cos φ (общий): {avg_cos:.4f}")

        # --- Таблица параметров по фазам (индивидуальные значения) ---
        if len(voltage_results_list) > 1:
            doc.add_heading('Параметры по фазам (индивидуальные)', level=1)
            cols = 8
            table = doc.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Фаза'
            hdr[1].text = 'f0, Гц'
            hdr[2].text = 'Urms, В'
            hdr[3].text = 'U откл., %'
            hdr[4].text = 'Ku, %'
            hdr[5].text = 'Статус Ku'
            hdr[6].text = 'Нарушений'
            hdr[7].text = 'Предупреждений'
            for idx, v_res in enumerate(voltage_results_list):
                row = table.add_row().cells
                phase = phase_labels[idx] if idx < len(phase_labels) else f'Фаза {chr(65+idx)}'
                row[0].text = phase
                row[1].text = f"{v_res['f0']:.3f}"
                row[2].text = f"{v_res['U_rms']:.1f}"
                row[3].text = f"{v_res['U_dev']:+.2f}"
                row[4].text = f"{v_res['Ku']:.2f}"
                row[5].text = v_res['Ku_status']
                row[6].text = str(len(v_res['violations']))
                row[7].text = str(len(v_res['warnings']))
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.runs:
                            run = p.runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1

        # --- Таблица параметров тока по фазам (если только ток с несколькими фазами) ---
        if only_current and len(current_results_list) > 1:
            doc.add_heading('Параметры тока по фазам', level=1)
            cols = 5
            table = doc.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Фаза'
            hdr[1].text = 'Irms, А'
            hdr[2].text = 'THDI, %'
            hdr[3].text = 'I_peak, А'
            hdr[4].text = 'Crest factor'
            for idx, i_res in enumerate(current_results_list):
                row = table.add_row().cells
                phase = phase_labels[idx] if idx < len(phase_labels) else f'Фаза {chr(65+idx)}'
                row[0].text = phase
                row[1].text = f"{i_res['I_rms']:.2f}"
                row[2].text = f"{i_res['THD_I']:.2f}"
                row[3].text = f"{i_res['I_peak']:.2f}"
                row[4].text = f"{i_res['crest_factor_I']:.3f}"
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.runs:
                            run = p.runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1

        # --- Суммарные параметры трёхфазной системы (если есть) ---
        if self.summary_data is not None:
            sd = self.summary_data
            doc.add_heading('СУММАРНЫЕ ПАРАМЕТРЫ ТРЁХФАЗНОЙ СИСТЕМЫ', level=1)
            doc.add_paragraph(f"Количество фаз: {sd['num_phases']}")
            doc.add_paragraph(f"Средняя частота: {sd['avg_freq']:.3f} Гц")
            doc.add_paragraph(f"Среднее действующее напряжение (RMS): {sd['avg_U_rms']:.1f} В")
            doc.add_paragraph(f"Среднее отклонение напряжения: {sd['avg_U_dev']:+.2f}%")
            doc.add_paragraph(f"Средний Ku (THD_U): {sd['avg_THD_U']:.2f}%")
            doc.add_paragraph(f"Средний THD_I: {sd['avg_THD_I']:.2f}%")
            doc.add_paragraph(f"Суммарная активная мощность P_total: {sd['P_total']:.2f} Вт ({sd['P_total']/1000:.3f} кВт)")
            doc.add_paragraph(f"Суммарная реактивная мощность Q_total: {sd['Q_total']:.2f} ВАр")
            doc.add_paragraph(f"Суммарная полная мощность S_total: {sd['S_total']:.2f} ВА")
            doc.add_paragraph(f"Общий коэффициент мощности: {sd['cos_phi_total']:.4f}")
            doc.add_paragraph(f"Всего нарушений по гармоникам: {sd['total_violations']}")
            doc.add_paragraph(f"Всего предупреждений: {sd['total_warnings']}")
            doc.add_paragraph(f"Общий статус Ku: {sd['ku_status_overall']}")
            doc.add_paragraph('')

        # ================================================================
        # 1. СВОДНЫЕ ТАБЛИЦЫ ДЛЯ РЕЖИМОВ "ТОЛЬКО НАПРЯЖЕНИЕ" И "ТОЛЬКО ТОК"
        # ================================================================

        # Режим "только напряжение" с несколькими фазами – сводная таблица
        if only_voltage and len(voltage_results_list) > 1:
            doc.add_heading('Сводная таблица гармоник напряжения по фазам', level=1)
            cols = 4 + len(voltage_results_list)
            table = doc.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = '№'
            hdr[1].text = 'Частота, Гц'
            col_idx = 2
            for phase in phase_labels:
                hdr[col_idx].text = f'Отн. ампл. {phase}, %'
                col_idx += 1
            hdr[col_idx].text = 'Предел ГОСТ, %'
            for k in range(1, 41):
                row = table.add_row().cells
                row[0].text = str(k)
                row[1].text = f"{voltage_results_list[0]['harm_freqs'][k-1]:.2f}"
                col_idx = 2
                for v_res in voltage_results_list:
                    rel = v_res['rel_amps'][k-1]
                    row[col_idx].text = f"{rel:.2f}"
                    col_idx += 1
                limit = voltage_results_list[0]['gost_limits'][k-1]
                row[col_idx].text = f"{limit:.2f}" if not np.isnan(limit) else '---'
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.runs:
                            run = p.runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1

        # Режим "только ток" с несколькими фазами – сводная таблица тока
        if only_current and len(current_results_list) > 1:
            doc.add_heading('Сводная таблица гармоник тока по фазам', level=1)
            cols = 3 + len(current_results_list)
            table = doc.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = '№'
            hdr[1].text = 'Частота, Гц'
            col_idx = 2
            for phase in phase_labels:
                hdr[col_idx].text = f'I_ампл {phase}, А'
                col_idx += 1
            for k in range(1, 41):
                row = table.add_row().cells
                row[0].text = str(k)
                freq = k * 50.0
                row[1].text = f"{freq:.2f}"
                col_idx = 2
                for i_res in current_results_list:
                    amp = i_res['harm_amps'][k-1]
                    row[col_idx].text = f"{amp:.4f}"
                    col_idx += 1
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.runs:
                            run = p.runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1

        # Режим "три фазы (3 файла)" – таблица сравнения относительных амплитуд
        if self.three_phase_mode and len(voltage_results_list) == 3:
            doc.add_heading('Сравнение относительных амплитуд гармоник по фазам', level=1)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = '№'
            hdr[1].text = 'Частота, Гц'
            hdr[2].text = 'Отн. ампл. Фаза A, %'
            hdr[3].text = 'Отн. ампл. Фаза B, %'
            hdr[4].text = 'Отн. ампл. Фаза C, %'
            hdr[5].text = 'Предел ГОСТ, %'
            for k in range(1, 41):
                row = table.add_row().cells
                row[0].text = str(k)
                row[1].text = f"{voltage_results_list[0]['harm_freqs'][k-1]:.2f}"
                for i, v_res in enumerate(voltage_results_list):
                    row[2+i].text = f"{v_res['rel_amps'][k-1]:.2f}"
                limit = voltage_results_list[0]['gost_limits'][k-1]
                row[5].text = f"{limit:.2f}" if not np.isnan(limit) else '---'
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.runs:
                            run = p.runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1

        # ================================================================
        # 2. ПОФАЗНЫЕ ТАБЛИЦЫ И ПАРАМЕТРЫ (для режимов с одной фазой или "напряжение+ток")
        # ================================================================

        # Если режим "только напряжение" с одной фазой или "напряжение+ток" – выводим пофазные таблицы
        if (only_voltage and len(voltage_results_list) == 1) or both:
            for idx, v_res in enumerate(voltage_results_list):
                phase = phase_labels[idx] if phase_labels else f'Фаза {chr(65+idx)}'
                doc.add_heading(f'Результаты для {phase}', level=1)
                doc.add_paragraph(f"Основная частота: {v_res['f0']:.3f} Гц")
                doc.add_paragraph(f"Отклонение частоты: {v_res['freq_dev']:+.3f} Гц ({v_res['freq_status']})")
                doc.add_paragraph(f"Действующее напряжение (RMS): {v_res['U_rms']:.1f} В")
                doc.add_paragraph(f"Отклонение напряжения: {v_res['U_dev']:+.2f}% (допустимо ±10%) - {v_res['U_status']}")
                doc.add_paragraph(f"Ku (THD): {v_res['Ku']:.2f}% (допустим {v_res['Ku_limit_95']:.1f}% для 95% времени, {v_res['Ku_limit_100']:.1f}% для 100%) - {v_res['Ku_status']}")
                doc.add_paragraph(f"Нарушений по гармоникам: {len(v_res['violations'])}")
                doc.add_paragraph(f"Предупреждений: {len(v_res['warnings'])}")

                doc.add_heading(f'Детальная таблица гармоник для {phase}', level=2)
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                headers = ['№', 'Частота, Гц', 'Амплитуда, В', 'Отн. ампл., %', 'Предел ГОСТ, %', 'Отклонение, %', 'Статус', 'Таблица ГОСТ']
                for i, h in enumerate(headers):
                    hdr[i].text = h
                for k in range(1, 41):
                    row = table.add_row().cells
                    row[0].text = str(k)
                    row[1].text = f"{v_res['harm_freqs'][k-1]:.2f}"
                    row[2].text = f"{v_res['harm_amps'][k-1]:.4f}"
                    row[3].text = f"{v_res['rel_amps'][k-1]:.2f}"
                    limit = v_res['gost_limits'][k-1]
                    row[4].text = f"{limit:.2f}" if not np.isnan(limit) else '---'
                    if k == 1:
                        row[5].text = '---'
                    else:
                        dev = v_res['rel_amps'][k-1] - limit if not np.isnan(limit) else np.nan
                        row[5].text = f"{dev:+.2f}" if not np.isnan(dev) else '---'
                    row[6].text = v_res['status_harm'][k-1]
                    table_str = v_res['gost_tables'][k-1]
                    if 'Таблица 1' in table_str:
                        table_str = 'Табл.1'
                    elif 'Таблица 2' in table_str:
                        table_str = 'Табл.2'
                    elif 'Таблица 3' in table_str:
                        table_str = 'Табл.3'
                    else:
                        table_str = '---'
                    row[7].text = table_str
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.runs:
                                run = p.runs[0]
                                run.font.size = Pt(10)
                                run.font.name = 'Times New Roman'
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.line_spacing = 1

                if both and idx < len(current_results_list):
                    i_res = current_results_list[idx]
                    doc.add_heading(f'Таблица гармоник тока для {phase}', level=2)
                    table_i = doc.add_table(rows=1, cols=4)
                    table_i.style = 'Table Grid'
                    hdr_i = table_i.rows[0].cells
                    hdr_i[0].text = '№'
                    hdr_i[1].text = 'Частота, Гц'
                    hdr_i[2].text = 'I_ампл, А'
                    hdr_i[3].text = 'Отн., %'
                    fund_I = i_res['harm_amps'][0]
                    for k in range(1, 41):
                        row_i = table_i.add_row().cells
                        row_i[0].text = str(k)
                        row_i[1].text = f"{v_res['harm_freqs'][k-1]:.2f}"
                        row_i[2].text = f"{i_res['harm_amps'][k-1]:.4f}"
                        rel = 100 * i_res['harm_amps'][k-1] / fund_I if fund_I > 0 else 0
                        row_i[3].text = f"{rel:.2f}"
                    for row in table_i.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.runs:
                                    run = p.runs[0]
                                    run.font.size = Pt(10)
                                    run.font.name = 'Times New Roman'
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.paragraph_format.line_spacing = 1

                    doc.add_heading(f'Параметры тока для {phase}', level=2)
                    doc.add_paragraph(f"Действующее значение тока Irms: {i_res['I_rms']:.2f} А")
                    doc.add_paragraph(f"THDI: {i_res['THD_I']:.2f}%")
                    if power_results_list and idx < len(power_results_list):
                        pr = power_results_list[idx]
                        doc.add_paragraph(f"Активная мощность P: {pr['P']:.2f} Вт ({pr['P']/1000:.3f} кВт)")
                        doc.add_paragraph(f"Реактивная мощность Q: {pr['Q']:.2f} ВАр")
                        doc.add_paragraph(f"Полная мощность S: {pr['S']:.2f} ВА")
                        doc.add_paragraph(f"Коэффициент мощности (общий): {pr['cos_phi_total']:.4f}")
                        doc.add_paragraph(f"cos φ (50 Гц): {pr['cos_phi_harm'][0]:.4f}")
                        doc.add_paragraph(f"Сдвиг фаз основной гармоники: {pr['delta_phase_deg'][0]:.1f}°")
                        PF_limit = 0.95
                        PF_status = 'СООТВЕТСТВУЕТ' if pr['cos_phi_total'] >= PF_limit else 'НЕ СООТВЕТСТВУЕТ'
                        doc.add_paragraph(f"PF (IEC 62040-3-2024): {pr['cos_phi_total']:.4f} (предел {PF_limit}) - {PF_status}")

        # --- Если режим "только ток" с одной фазой – выводим параметры тока ---
        if only_current and len(current_results_list) == 1:
            i_res = current_results_list[0]
            phase = phase_labels[0] if phase_labels else 'Ток'
            doc.add_heading(f'Результаты для {phase}', level=1)
            doc.add_heading(f'Параметры тока', level=2)
            doc.add_paragraph(f"Действующее значение тока Irms: {i_res['I_rms']:.2f} А")
            doc.add_paragraph(f"THDI: {i_res['THD_I']:.2f}%")
            doc.add_paragraph(f"Максимальный мгновенный ток: {np.max(i_res['signal']):.3f} А")
            doc.add_paragraph(f"Минимальный мгновенный ток: {np.min(i_res['signal']):.3f} А")
            doc.add_paragraph(f"Пиковый ток (peak): {i_res['I_peak']:.2f} А")
            doc.add_paragraph(f"Коэффициент амплитуды тока (peak/RMS): {i_res['crest_factor_I']:.3f}:1 (для чистого синуса √2 ≈ 1.414)")
            doc.add_heading(f'Таблица гармоник тока для {phase}', level=2)
            table_i = doc.add_table(rows=1, cols=4)
            table_i.style = 'Table Grid'
            hdr_i = table_i.rows[0].cells
            hdr_i[0].text = '№'
            hdr_i[1].text = 'Частота, Гц'
            hdr_i[2].text = 'I_ампл, А'
            hdr_i[3].text = 'Отн., %'
            fund_I = i_res['harm_amps'][0]
            for k in range(1, 41):
                row_i = table_i.add_row().cells
                row_i[0].text = str(k)
                row_i[1].text = f"{k*50:.2f}"
                row_i[2].text = f"{i_res['harm_amps'][k-1]:.4f}"
                rel = 100 * i_res['harm_amps'][k-1] / fund_I if fund_I > 0 else 0
                row_i[3].text = f"{rel:.2f}"
            for row in table_i.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.runs:
                            run = p.runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1

        # --- Итоговый вывод и рекомендации (для всех режимов) ---
        if voltage_results_list or current_results_list:
            total_viol = 0
            total_warn = 0
            if voltage_results_list:
                total_viol += sum(len(v['violations']) for v in voltage_results_list)
                total_warn += sum(len(v['warnings']) for v in voltage_results_list)
                ku_viol = sum(1 for v in voltage_results_list if v['Ku_status'] != 'СООТВЕТСТВУЕТ')
            else:
                ku_viol = 0
            total_issues = total_viol + ku_viol

            doc.add_heading('ИТОГОВЫЙ ВЫВОД:', level=1)
            if total_issues == 0:
                doc.add_paragraph('Качество электроэнергии ОТЛИЧНОЕ')
                doc.add_paragraph('- Полное соответствие ГОСТ 32144-2013')
                doc.add_paragraph('- Значительный запас по всем параметрам')
            elif total_issues <= 3:
                doc.add_paragraph('Качество электроэнергии УДОВЛЕТВОРИТЕЛЬНОЕ')
                doc.add_paragraph('- Соответствует ГОСТ 32144-2013 для 100% времени')
                if total_warn > 0:
                    doc.add_paragraph(f'- Имеются предупреждения ({total_warn})')
            else:
                doc.add_paragraph('Качество электроэнергии НЕУДОВЛЕТВОРИТЕЛЬНОЕ')
                doc.add_paragraph('- Обнаружены нарушения ГОСТ 32144-2013')
                doc.add_paragraph('- Требуется корректирующее воздействие')

            doc.add_heading('РЕКОМЕНДАЦИИ:', level=1)
            if total_viol > 0:
                doc.add_paragraph('1. Установка фильтров гармоник для снижения уровня высших гармоник.')
            if voltage_results_list and any(abs(v['U_dev']) > 10 for v in voltage_results_list):
                doc.add_paragraph('2. Проверить работу регулятора напряжения.')
            if ku_viol > 0:
                doc.add_paragraph('3. Снизить коэффициент искажения синусоидальности Ku.')
            doc.add_paragraph('')  # пустая строка

        # --- Дополнительные параметры сигналов ---
        if voltage_results_list:
            v_res = voltage_results_list[0]
            doc.add_heading('--- ДОПОЛНИТЕЛЬНЫЕ ПАРАМЕТРЫ СИГНАЛОВ ---', level=1)
            doc.add_paragraph(f"Максимальное мгновенное напряжение: {v_res['U_max']:.2f} В")
            doc.add_paragraph(f"Минимальное мгновенное напряжение: {v_res['U_min']:.2f} В")
            doc.add_paragraph(f"Пиковое напряжение (peak): {v_res['U_peak']:.2f} В")
            doc.add_paragraph(f"Коэффициент амплитуды напряжения (peak/RMS): {v_res['crest_factor_U']:.3f}:1 (для чистого синуса √2 ≈ 1.414)")

        if current_results_list:
            i_res = current_results_list[0]
            doc.add_paragraph(f"Максимальный мгновенный ток: {np.max(i_res['signal']):.3f} А")
            doc.add_paragraph(f"Минимальный мгновенный ток: {np.min(i_res['signal']):.3f} А")
            doc.add_paragraph(f"Пиковый ток (peak): {i_res['I_peak']:.2f} А")
            doc.add_paragraph(f"Коэффициент амплитуды тока (peak/RMS): {i_res['crest_factor_I']:.3f}:1 (для чистого синуса √2 ≈ 1.414)")

        # --- Если есть ток и напряжение, добавляем таблицу U/I/Δφ/cosφ и проверку по IEC ---
        if both and power_results_list:
            pr = power_results_list[0]
            doc.add_heading('--- ТАБЛИЦА ГАРМОНИК (1..40) ---', level=1)
            table_ui = doc.add_table(rows=1, cols=6)
            table_ui.style = 'Table Grid'
            hdr_ui = table_ui.rows[0].cells
            hdr_ui[0].text = '№'
            hdr_ui[1].text = 'Частота, Гц'
            hdr_ui[2].text = 'U_ампл, В'
            hdr_ui[3].text = 'I_ампл, А'
            hdr_ui[4].text = 'Δφ, °'
            hdr_ui[5].text = 'cos φ'
            for k in range(1, 41):
                row = table_ui.add_row().cells
                row[0].text = str(k)
                freq = voltage_results_list[0]['harm_freqs'][k-1]
                row[1].text = f"{freq:.2f}"
                row[2].text = f"{voltage_results_list[0]['harm_amps'][k-1]:.4f}"
                row[3].text = f"{current_results_list[0]['harm_amps'][k-1]:.4f}"
                row[4].text = f"{pr['delta_phase_deg'][k-1]:.1f}"
                row[5].text = f"{pr['cos_phi_harm'][k-1]:.4f}"
            for row in table_ui.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.runs:
                            run = p.runs[0]
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1

            doc.add_heading('--- ПРОВЕРКА ПО ГОСТ IEC 62040-3-2024 ---', level=1)
            i_res = current_results_list[0]
            v_res = voltage_results_list[0]
            THDi_limit = 5
            THDi_status = 'СООТВЕТСТВУЕТ' if i_res['THD_I'] < THDi_limit else 'НЕ СООТВЕТСТВУЕТ'
            doc.add_paragraph(f"THD тока (THDi): {i_res['THD_I']:.2f}% (предел {THDi_limit}%) - {THDi_status}")

            THDu_limit = 8
            THDu_status = 'СООТВЕТСТВУЕТ' if v_res['Ku'] < THDu_limit else 'НЕ СООТВЕТСТВУЕТ'
            doc.add_paragraph(f"THD напряжения (THDu) для нелинейной нагрузки: {v_res['Ku']:.2f}% (предел {THDu_limit}%) - {THDu_status}")

            PF_limit = 0.95
            PF_status = 'СООТВЕТСТВУЕТ' if pr['cos_phi_total'] >= PF_limit else 'НЕ СООТВЕТСТВУЕТ'
            doc.add_paragraph(f"Коэффициент мощности (PF): {pr['cos_phi_total']:.4f} (предел {PF_limit}) - {PF_status}")

        # --- Добавление графиков ---
        if image_paths:
            doc.add_heading('Графики', level=1)
            for img_path in image_paths:
                if 'long_term' not in img_path:
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph()

        # --- Длительный анализ (если есть) ---
        if self.long_term_results is not None and 'error' not in self.long_term_results:
            res = self.long_term_results
            doc.add_heading('Длительный анализ', level=1)
            doc.add_paragraph(f"Количество периодов: {res['num_periods']}")
            doc.add_paragraph(f"Минимальное RMS: {res['U_min_rms']:.2f} В")
            doc.add_paragraph(f"Максимальное RMS: {res['U_max_rms']:.2f} В")
            doc.add_paragraph(f"Среднее RMS: {res['U_mean_rms']:.2f} В")
            doc.add_paragraph(f"Стандартное отклонение RMS: {res['U_std_rms']:.3f} В")
            doc.add_paragraph(f"Максимальный пик: {res['peak_abs_max']:.2f} В")
            doc.add_paragraph(f"Дельта (пик - мин.RMS): {res['delta']:.2f} В")
            doc.add_paragraph(f"Количество нарушений (±10%): {res['violations']}")

            rms_values = res.get('rms_values', [])
            if rms_values:
                doc.add_heading('Таблица RMS по периодам', level=2)
                total = len(rms_values)
                step = self._get_rms_step(total)
                doc.add_paragraph(f"Показаны значения каждые {step} периодов (всего {total} периодов)")
                table_rms = doc.add_table(rows=1, cols=2)
                table_rms.style = 'Table Grid'
                hdr = table_rms.rows[0].cells
                hdr[0].text = 'Номер периода'
                hdr[1].text = 'RMS, В'
                for i in range(0, total, step):
                    row = table_rms.add_row().cells
                    row[0].text = str(i+1)
                    row[1].text = f"{rms_values[i]:.2f}"
                if (total - 1) % step != 0:
                    row = table_rms.add_row().cells
                    row[0].text = str(total)
                    row[1].text = f"{rms_values[-1]:.2f}"
                for row in table_rms.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.runs:
                                run = p.runs[0]
                                run.font.size = Pt(10)
                                run.font.name = 'Times New Roman'
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.line_spacing = 1

            rms_img = os.path.join(self.output_folder, f"long_term_rms_{self.base_filename}.png")
            if os.path.exists(rms_img):
                doc.add_picture(rms_img, width=Inches(6))
                doc.add_paragraph()

        doc.save(self.word_path)
        return self.word_path
# ----------------------------------------------------------------------
# Диалог настроек (отладка) – только масштабы U/I, инверсия, U_nom
# ----------------------------------------------------------------------
class DebugDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Настройки анализа (отладка)")
        self.setMinimumWidth(350)
        layout = QVBoxLayout(self)

        h1 = QHBoxLayout()
        h1.addWidget(QLabel("Масштаб напряжения (x):"))
        self.scale_U = QDoubleSpinBox()
        self.scale_U.setRange(0.01, 10000)
        self.scale_U.setValue(1.0)
        self.scale_U.setDecimals(3)
        h1.addWidget(self.scale_U)
        layout.addLayout(h1)

        h2 = QHBoxLayout()
        h2.addWidget(QLabel("Масштаб тока (x):"))
        self.scale_I = QDoubleSpinBox()
        self.scale_I.setRange(0.01, 10000)
        self.scale_I.setValue(1694.0)
        self.scale_I.setDecimals(3)
        h2.addWidget(self.scale_I)
        layout.addLayout(h2)

        self.invert_I = QCheckBox("Инвертировать ток")
        layout.addWidget(self.invert_I)

        h3 = QHBoxLayout()
        h3.addWidget(QLabel("Номинальное напряжение (В):"))
        self.U_nom = QDoubleSpinBox()
        self.U_nom.setRange(1, 100000)
        self.U_nom.setValue(230.0)
        self.U_nom.setDecimals(1)
        h3.addWidget(self.U_nom)
        layout.addLayout(h3)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def get_values(self):
        return {
            'scale_U': self.scale_U.value(),
            'scale_I': self.scale_I.value(),
            'invert_I': self.invert_I.isChecked(),
            'U_nom': self.U_nom.value()
        }

    def set_values(self, scale_U, scale_I, invert_I, U_nom):
        self.scale_U.setValue(scale_U)
        self.scale_I.setValue(scale_I)
        self.invert_I.setChecked(invert_I)
        self.U_nom.setValue(U_nom)

# ----------------------------------------------------------------------
# Основное приложение (без бинарных файлов)
# ----------------------------------------------------------------------
class MainApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Анализатор гармоник по ГОСТ 32144-2013")
        self.setMinimumSize(800, 600)
        self.output_folder = None
        self.file_paths = []
        self.long_term_results = None
        # Параметры, доступные только через кнопку "Настройки (отладка)"
        self.scale_U = 1.0
        self.scale_I = 1694.0
        self.invert_I = False
        self.U_nom = 230.0
        self.initUI()
        logger.set_gui(self.log_text)

    def initUI(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # Выбор файла
        file_layout = QHBoxLayout()
        self.file_path_edit = QLineEdit()
        self.file_path_edit.setReadOnly(True)
        browse_btn = QPushButton("Выбрать файл")
        browse_btn.clicked.connect(self.browse_file)
        file_layout.addWidget(QLabel("Файл:"))
        file_layout.addWidget(self.file_path_edit)
        file_layout.addWidget(browse_btn)
        layout.addLayout(file_layout)

        # Формат
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel("Формат осциллографа:"))
        self.format_combo = QComboBox()
        self.format_combo.addItems(["Rigol old", "Rigol DHO814", "AKIP"])
        format_layout.addWidget(self.format_combo)
        layout.addLayout(format_layout)

        # Режим анализа
        mode_group = QGroupBox("Режим анализа")
        mode_layout = QHBoxLayout()
        self.mode_voltage = QRadioButton("Только напряжение")
        self.mode_current = QRadioButton("Только ток")
        self.mode_both = QRadioButton("Напряжение и ток")
        self.mode_voltage.setChecked(True)
        self.mode_three_phases = QRadioButton("Три фазы (3 файла)")
        mode_layout.addWidget(self.mode_three_phases)
        mode_layout.addWidget(self.mode_voltage)
        mode_layout.addWidget(self.mode_current)
        mode_layout.addWidget(self.mode_both)
        mode_group.setLayout(mode_layout)
        layout.addWidget(mode_group)

        # Длительный анализ
        long_group = QGroupBox("Длительный анализ (для больших файлов)")
        long_layout = QHBoxLayout()
        self.long_analysis = QCheckBox("Включить длительный анализ")
        self.long_analysis.setChecked(False)
        long_layout.addWidget(self.long_analysis)
        long_layout.addWidget(QLabel("Номинал U, В:"))
        self.long_U_nom = QDoubleSpinBox()
        self.long_U_nom.setRange(1, 100000)
        self.long_U_nom.setValue(230.0)
        self.long_U_nom.setDecimals(1)
        long_layout.addWidget(self.long_U_nom)
        long_layout.addWidget(QLabel("Допуск, %:"))
        self.long_tolerance = QDoubleSpinBox()
        self.long_tolerance.setRange(0.1, 50)
        self.long_tolerance.setValue(10.0)
        self.long_tolerance.setDecimals(1)
        long_layout.addWidget(self.long_tolerance)
        long_group.setLayout(long_layout)
        layout.addWidget(long_group)

        # Назначение каналов
        channels_group = QGroupBox("Назначение каналов")
        channels_layout = QGridLayout()
        labels = ["Канал 1", "Канал 2", "Канал 3", "Канал 4"]
        self.channel_combo = []
        for i in range(4):
            channels_layout.addWidget(QLabel(labels[i]), i, 0)
            combo = QComboBox()
            combo.addItems(["Не используется", "Напряжение", "Ток"])
            channels_layout.addWidget(combo, i, 1)
            self.channel_combo.append(combo)
        channels_group.setLayout(channels_layout)
        layout.addWidget(channels_group)

        # Папка результатов
        folder_layout = QHBoxLayout()
        folder_layout.addWidget(QLabel("Папка для сохранения:"))
        self.folder_path_edit = QLineEdit()
        self.folder_path_edit.setReadOnly(True)
        folder_btn = QPushButton("Выбрать папку")
        folder_btn.clicked.connect(self.browse_output_folder)
        folder_layout.addWidget(self.folder_path_edit)
        folder_layout.addWidget(folder_btn)
        layout.addLayout(folder_layout)

        # Кнопки
        btn_layout = QHBoxLayout()
        self.run_btn = QPushButton("Запустить анализ")
        self.run_btn.clicked.connect(self.run_analysis)
        btn_layout.addWidget(self.run_btn)

        self.debug_btn = QPushButton("Настройки (отладка)")
        self.debug_btn.clicked.connect(self.open_debug_dialog)
        btn_layout.addWidget(self.debug_btn)
        layout.addLayout(btn_layout)

        self.progress = QProgressBar()
        self.progress.setValue(0)
        layout.addWidget(self.progress)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)

        self.debug_dialog = None

    def browse_file(self):
        if self.mode_three_phases.isChecked():
            file_paths, _ = QFileDialog.getOpenFileNames(
                self,
                "Выберите 3 CSV-файла (фазы A, B, C)",
                "",
                "CSV files (*.csv);;All files (*.*)"
            )
            if file_paths:
                if len(file_paths) != 3:
                    QMessageBox.warning(self, "Ошибка", "Необходимо выбрать ровно 3 файла.")
                    return
                self.file_path_edit.setText("; ".join(file_paths))
                self.file_paths = file_paths
        else:
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Выберите файл данных",
                "",
                "CSV files (*.csv);;All files (*.*)"
            )
            if file_path:
                self.file_path_edit.setText(file_path)
                self.file_paths = [file_path]

    def browse_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения результатов")
        if folder:
            if not os.path.exists(folder):
                try:
                    os.makedirs(folder, exist_ok=True)
                except Exception as e:
                    QMessageBox.warning(self, "Ошибка", f"Не удалось создать папку: {e}")
                    return
            self.folder_path_edit.setText(folder)
            self.output_folder = folder
            log_path = os.path.join(folder, "log.txt")
            logger.set_file(log_path)

    def open_debug_dialog(self):
        if self.debug_dialog is None:
            self.debug_dialog = DebugDialog(self)
        self.debug_dialog.set_values(self.scale_U, self.scale_I, self.invert_I, self.U_nom)
        if self.debug_dialog.exec_() == QDialog.Accepted:
            vals = self.debug_dialog.get_values()
            self.scale_U = vals['scale_U']
            self.scale_I = vals['scale_I']
            self.invert_I = vals['invert_I']
            self.U_nom = vals['U_nom']
            logger.log(f"Настройки обновлены: U_scale={self.scale_U}, I_scale={self.scale_I}, инверсия={self.invert_I}, U_nom={self.U_nom} В")

    def run_analysis(self):
        file_paths = self.file_paths
        if not file_paths:
            QMessageBox.warning(self, "Ошибка", "Выберите файл(ы).")
            return
        for fp in file_paths:
            if not os.path.exists(fp):
                QMessageBox.warning(self, "Ошибка", f"Файл не существует: {fp}")
                return
        if not self.output_folder:
            QMessageBox.warning(self, "Ошибка", "Выберите папку для сохранения результатов.")
            return

        self.progress.setValue(0)
        logger.log("Начало анализа")

        format_type = self.format_combo.currentText()
        mode_three_phases = self.mode_three_phases.isChecked()
        mode = "voltage" if self.mode_voltage.isChecked() else "current" if self.mode_current.isChecked() else "both"

        # Для трёхфазного режима принудительно устанавливаем режим "both"
        if mode_three_phases:
            mode = "both"

        channel_types = []
        for i in range(4):
            ch_type = self.channel_combo[i].currentText()
            channel_types.append(ch_type)

        # В трёхфазном режиме ожидаем ровно один канал напряжения и один канал тока в каждом файле
        if mode_three_phases and (channel_types.count("Напряжение") != 1 or channel_types.count("Ток") != 1):
            QMessageBox.warning(self, "Ошибка", "Для трёхфазного режима необходимо выбрать ровно один канал как напряжение и один как ток (они будут применены ко всем трём файлам).")
            return

        if mode == "voltage" and channel_types.count("Напряжение") == 0:
            QMessageBox.warning(self, "Ошибка", "Для режима 'Только напряжение' необходимо выбрать хотя бы один канал как напряжение.")
            return
        if mode == "current" and channel_types.count("Ток") == 0:
            QMessageBox.warning(self, "Ошибка", "Для режима 'Только ток' необходимо выбрать хотя бы один канал как ток.")
            return

        voltage_results_list = []
        current_results_list = []
        power_results_list = []
        long_res = None

        # ---- Цикл по файлам ----
        for file_idx, file_path in enumerate(file_paths):
            try:
                logger.log(f"Чтение CSV-файла {file_idx+1}: {os.path.basename(file_path)}")
                t, data = read_oscilloscope_file(file_path, format_type)
                N, M = data.shape
                logger.log(f"Прочитано {N} отсчётов, {M} каналов.")
            except Exception as e:
                logger.log(f"Ошибка чтения файла {file_path}: {e}")
                QMessageBox.critical(self, "Ошибка", f"Не удалось прочитать файл {file_path}: {e}")
                return

            # Извлечение сигналов согласно назначению каналов
            voltage_signals = []
            current_signals = []
            for i, ch_type in enumerate(channel_types):
                if i < M:
                    if ch_type == "Напряжение":
                        voltage_signals.append(data[:, i] * self.scale_U)
                    elif ch_type == "Ток":
                        sig = data[:, i] * self.scale_I
                        if self.invert_I:
                            sig = -sig
                        current_signals.append(sig)

            if len(voltage_signals) + len(current_signals) == 0:
                QMessageBox.warning(self, "Ошибка", f"В файле {os.path.basename(file_path)} не выбрано ни одного канала для анализа.")
                return

            fs = 1.0 / np.mean(np.diff(t))
            logger.log(f"Частота дискретизации (расчётная): {fs:.2f} Гц")

            gost = GOSTAnalyzer(voltage_level=1, U_nom=self.U_nom)

                        # ---- Анализ напряжения (все назначенные каналы) ----
            if voltage_signals:
                for v_sig in voltage_signals:
                    try:
                        v_res = gost.analyze_voltage(v_sig, fs)
                        voltage_results_list.append(v_res)
                        logger.log(f"Напряжение (файл {file_idx+1}, канал {len(voltage_results_list)}): f0={v_res['f0']:.3f} Гц, U_rms={v_res['U_rms']:.1f} В, Ku={v_res['Ku']:.2f}%")
                    except Exception as e:
                        logger.log(f"Ошибка анализа напряжения в файле {file_path}: {e}")
                        QMessageBox.critical(self, "Ошибка", f"Анализ напряжения не удался: {e}")
                        return
            # ---- Анализ тока (все назначенные каналы) ----
            if current_signals:
                # Если есть напряжения, используем соответствующие частоты
                if voltage_results_list:
                    # Если число токов совпадает с числом напряжений – используем попарно
                    if len(current_signals) == len(voltage_results_list):
                        for idx, i_sig in enumerate(current_signals):
                            f0 = voltage_results_list[idx]['f0']
                            try:
                                i_res = gost.analyze_current(i_sig, fs, f0)
                                current_results_list.append(i_res)
                                logger.log(f"Ток (файл {file_idx+1}, канал {idx+1}): I_rms={i_res['I_rms']:.2f} А, THD_I={i_res['THD_I']:.2f}%")
                            except Exception as e:
                                logger.log(f"Ошибка анализа тока в файле {file_path}: {e}")
                                QMessageBox.critical(self, "Ошибка", f"Анализ тока не удался: {e}")
                                return
                    else:
                        # Если число токов не совпадает, используем первую частоту для всех
                        f0 = voltage_results_list[0]['f0']
                        for idx, i_sig in enumerate(current_signals):
                            try:
                                i_res = gost.analyze_current(i_sig, fs, f0)
                                current_results_list.append(i_res)
                                logger.log(f"Ток (файл {file_idx+1}, канал {idx+1}): I_rms={i_res['I_rms']:.2f} А, THD_I={i_res['THD_I']:.2f}%")
                            except Exception as e:
                                logger.log(f"Ошибка анализа тока в файле {file_path}: {e}")
                                QMessageBox.critical(self, "Ошибка", f"Анализ тока не удался: {e}")
                                return
                else:
                    # Если нет напряжения, оцениваем частоту по первому току
                    tmp = gost.analyze_voltage(current_signals[0], fs)
                    f0 = tmp['f0']
                    for idx, i_sig in enumerate(current_signals):
                        try:
                            i_res = gost.analyze_current(i_sig, fs, f0)
                            current_results_list.append(i_res)
                            logger.log(f"Ток (файл {file_idx+1}, канал {idx+1}): I_rms={i_res['I_rms']:.2f} А, THD_I={i_res['THD_I']:.2f}%")
                        except Exception as e:
                            logger.log(f"Ошибка анализа тока в файле {file_path}: {e}")
                            QMessageBox.critical(self, "Ошибка", f"Анализ тока не удался: {e}")
                            return

            # ---- Расчёт мощности для каждой пары (если есть и напряжение, и ток) ----
            if voltage_results_list and current_results_list:
                # Берём пары по минимуму (если количество не совпадает)
                num_pairs = min(len(voltage_results_list), len(current_results_list))
                # Но для режима "both" ожидается, что они равны
                # Используем все пары только если они равны
                if len(voltage_results_list) == len(current_results_list):
                    for idx in range(len(voltage_results_list)):
                        v_res = voltage_results_list[idx]
                        i_res = current_results_list[idx]
                        power = PowerAnalyzer.compute_power(
                            v_res['harm_amps'], i_res['harm_amps'],
                            v_res['harm_phases'], i_res['harm_phases'],
                            v_res['U_rms'], i_res['I_rms']
                        )
                        power_results_list.append(power)
                        logger.log(f"Мощность (файл {file_idx+1}, пара {idx+1}): P={power['P']:.2f} Вт, Q={power['Q']:.2f} ВАр, cosφ={power['cos_phi_total']:.4f}")
                else:
                    # Если не совпадает – предупреждение и только первые пары
                    logger.log(f"Предупреждение: количество каналов напряжения ({len(voltage_results_list)}) и тока ({len(current_results_list)}) не совпадает. Используются только первые {num_pairs} пар.")
                    for idx in range(num_pairs):
                        v_res = voltage_results_list[idx]
                        i_res = current_results_list[idx]
                        power = PowerAnalyzer.compute_power(
                            v_res['harm_amps'], i_res['harm_amps'],
                            v_res['harm_phases'], i_res['harm_phases'],
                            v_res['U_rms'], i_res['I_rms']
                        )
                        power_results_list.append(power)
                        logger.log(f"Мощность (файл {file_idx+1}, пара {idx+1}): P={power['P']:.2f} Вт, Q={power['Q']:.2f} ВАр, cosφ={power['cos_phi_total']:.4f}")

            # Длительный анализ (только для первого файла)
            if self.long_analysis.isChecked() and file_idx == 0 and voltage_results_list:
                logger.log("Запуск длительного анализа...")
                v_sig = voltage_results_list[0]['signal']
                fs = voltage_results_list[0]['Fs']
                U_nom = self.long_U_nom.value()
                tolerance = self.long_tolerance.value()
                long_res = LongTermAnalyzer.analyze(v_sig, fs, U_nom, freq_estimate=voltage_results_list[0]['f0'], rms_tolerance=tolerance)
                if 'error' in long_res:
                    logger.log(f"Ошибка длительного анализа: {long_res['error']}")
                    long_res = None
                else:
                    logger.log(f"Длительный анализ: периодов={long_res['num_periods']}, U_min_rms={long_res['U_min_rms']:.1f} В, "
                               f"пик={long_res['peak_abs_max']:.1f} В, дельта={long_res['delta']:.1f} В, "
                               f"нарушений={long_res['violations']}")

        self.long_term_results = long_res
        self.progress.setValue(60)

        # ---- Вычисление суммарных параметров для трёхфазной системы ----
        summary_data = None
        if mode_three_phases and len(voltage_results_list) == 3 and len(current_results_list) == 3 and len(power_results_list) == 3:
            P_total = sum([p['P'] for p in power_results_list])
            Q_total = sum([p['Q'] for p in power_results_list])
            S_total = np.sqrt(P_total**2 + Q_total**2)
            cos_phi_total = P_total / S_total if S_total > 0 else 0

            avg_THD_U = np.mean([v['Ku'] for v in voltage_results_list])
            avg_THD_I = np.mean([i['THD_I'] for i in current_results_list])
            total_violations = sum([len(v['violations']) for v in voltage_results_list])
            total_warnings = sum([len(v['warnings']) for v in voltage_results_list])
            ku_statuses = [v['Ku_status'] for v in voltage_results_list]
            ku_status_overall = 'СООТВЕТСТВУЕТ' if all(s == 'СООТВЕТСТВУЕТ' for s in ku_statuses) else 'НЕ СООТВЕТСТВУЕТ'
            avg_U_dev = np.mean([v['U_dev'] for v in voltage_results_list])
            avg_U_rms = np.mean([v['U_rms'] for v in voltage_results_list])
            avg_freq = np.mean([v['f0'] for v in voltage_results_list])

            summary_data = {
                'P_total': P_total,
                'Q_total': Q_total,
                'S_total': S_total,
                'cos_phi_total': cos_phi_total,
                'avg_THD_U': avg_THD_U,
                'avg_THD_I': avg_THD_I,
                'total_violations': total_violations,
                'total_warnings': total_warnings,
                'ku_status_overall': ku_status_overall,
                'avg_U_dev': avg_U_dev,
                'avg_U_rms': avg_U_rms,
                'avg_freq': avg_freq,
                'num_phases': len(voltage_results_list)
            }
            logger.log(f"Суммарные параметры: P_total={P_total:.2f} Вт, Q_total={Q_total:.2f} ВАр, cosφ={cos_phi_total:.4f}")

        self.progress.setValue(70)

        # ---- Генерация отчётов ----
        # Базовое имя для отчётов – имя первого файла
        base_name = os.path.splitext(os.path.basename(file_paths[0]))[0]
        report_gen = ReportGenerator(self.output_folder, base_name)
        report_gen.long_term_results = self.long_term_results
        if summary_data is not None:
            report_gen.set_summary_data(summary_data)

        # Метки фаз
        if mode_three_phases and len(voltage_results_list) == 3:
            phase_labels = ['Фаза A', 'Фаза B', 'Фаза C']
        else:
            phase_labels = [f'Фаза {chr(65+i)}' for i in range(len(voltage_results_list))]

        logger.log("Генерация графиков...")
        try:
            image_paths = report_gen.generate_full_report(
                voltage_results_list,
                current_results_list if current_results_list else None,
                power_results_list if power_results_list else None,
                phase_labels
            )
            if current_results_list:
                for idx, (v_res, i_res) in enumerate(zip(voltage_results_list, current_results_list)):
                    img = report_gen.generate_vector_diagram(v_res, i_res, phase_labels[idx])
                    image_paths.append(img)
            logger.log(f"Создано {len(image_paths)} графиков")
        except Exception as e:
            logger.log(f"Ошибка генерации графиков: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Ошибка", f"Ошибка при создании графиков: {e}")
            return

        self.progress.setValue(85)

        logger.log("Формирование Excel-файла...")
        try:
            report_gen.generate_excel(voltage_results_list, current_results_list if current_results_list else None,
                                      power_results_list if power_results_list else None, phase_labels)
        except Exception as e:
            logger.log(f"Ошибка Excel: {e}")

        logger.log("Формирование Word-файла...")
        try:
            report_gen.generate_word(voltage_results_list, current_results_list if current_results_list else None,
                                     power_results_list if power_results_list else None, phase_labels, image_paths)
        except Exception as e:
            logger.log(f"Ошибка Word: {e}")

        self.progress.setValue(100)
        logger.log(f"Готово! Результаты сохранены в {self.output_folder}")
        QMessageBox.information(self, "Завершено", f"Анализ завершён. Результаты в папке:\n{self.output_folder}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainApp()
    window.show()
    sys.exit(app.exec_())